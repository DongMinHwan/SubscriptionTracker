#!/usr/bin/env python3
"""2026-08-22 심사 제출 오류 .docx 생성."""

import os
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from build_docx import (
    add_body,
    add_bullet,
    add_heading_styled,
    add_number,
    configure_document,
    make_table,
)
from docx import Document

BASE = "/Users/hwangdongmin/Projects/NewProject_1"
OUT_PATH = f"{BASE}/출시/2026-08-22_심사제출_오류.docx"
KV = [3.6, 13.0]


def build():
    doc = Document()
    configure_document(doc)

    add_heading_styled(doc, "내 구독 트래커 — 심사 제출 때 난 오류", 1)

    make_table(
        doc,
        ["항목", "내용"],
        [
            ["작성일", "2026년 8월 22일"],
            ["언제", "버전 페이지를 채운 뒤 처음 Add for Review를 눌렀을 때, 그리고 그 전후"],
            [
                "한 줄",
                "앱 코드가 틀린 게 아니라 Connect 필수 칸이 비어 있거나, "
                "칸 위치를 잘못 연 경우가 대부분이었다.",
            ],
            [
                "짝 문서",
                "출시/2026-08-22_심사제출과정.docx 에 올바른 순서와 넣은 값이 있다.",
            ],
        ],
        KV,
    )

    add_heading_styled(doc, "1. 첫 Add for Review — 빨간 상자 두 개", 2)
    add_body(
        doc,
        "시각은 2026년 8월 22일 11:45 무렵. "
        "화면 제목은 iOS App Version 1.0.0. "
        "위쪽 분홍 상자 제목은 Unable to Add for Review, "
        "바로 아래 상자는 You have one or more errors on this page.",
    )

    add_heading_styled(doc, "1.1 Unable to Add for Review", 3)
    add_body(
        doc,
        "앱 공통 정보가 비어 있을 때 나온다. "
        "왼쪽 사이드바의 다른 메뉴로 가서 채운 뒤 다시 누르면 된다.",
    )
    make_table(
        doc,
        ["화면에 나온 뜻", "어디를 열었는가", "원인", "어떻게 고쳤는가"],
        [
            [
                "Primary Category를 고르라",
                "App Information → Category",
                "주 카테고리가 비어 있음",
                "Primary = Finance. Secondary는 비움",
            ],
            [
                "Content Rights Information을 설정하라",
                "App Information → Content Rights → Edit",
                "콘텐츠 권리 설문이 비어 있음",
                "타사 콘텐츠 없음 (No third-party content)",
            ],
            [
                "Age Ratings 질문에 답하라",
                "App Information → Age Ratings → Set Up",
                "연령 설문이 비어 있음",
                "해당 문항 전부 NONE/NO. 결과 4+",
            ],
            [
                "Copyright information을 넣으라",
                "버전 페이지 Copyright 칸",
                "저작권 한 줄이 비어 있음. "
                "App Information이 아니라 버전 페이지에 있다",
                "2026 내 구독 트래커. 실명은 넣지 않음",
            ],
            [
                "Pricing에서 price tier를 고르라",
                "Pricing and Availability → Add Pricing",
                "가격이 한 번도 설정되지 않음",
                "기준 나라 한국(KRW), 가격 ₩0.00, 175개국",
            ],
        ],
        [3.8, 4.2, 4.0, 4.6],
    )

    add_heading_styled(doc, "1.2 You have one or more errors on this page", 3)
    add_body(
        doc,
        "지금 보고 있는 버전 페이지의 App Review Information이 비어 있을 때 나온다. "
        "Show Details를 누르면 칸 이름이 보인다.",
    )
    make_table(
        doc,
        ["칸", "원인", "어떻게 고쳤는가"],
        [
            ["First name is required", "심사 연락처 이름이 비어 있음", "DONGMIN"],
            ["Last name is required", "성이 비어 있음", "HWANG"],
            ["Email is required", "메일이 비어 있음", "hdmin8712@gmail.com"],
            [
                "Phone number is required",
                "전화가 비어 있음",
                "받을 수 있는 번호, +82 형식",
            ],
        ],
        [5.0, 5.6, 6.0],
    )
    add_body(
        doc,
        "이 네 칸은 스토어 제품 페이지가 아니라 심사자에게만 간다. "
        "Copyright 줄과는 다르다.",
        before=8,
    )

    add_heading_styled(doc, "2. 제출 버튼을 누르기 전에 막혔던 것", 2)
    add_body(
        doc,
        "빨간 상자 전에 이미 몇 번 멈췄다. "
        "Connect가 오류라고 크게 쓰지 않아도, 칸을 못 찾거나 저장이 안 되면 같은 일이다.",
    )

    add_heading_styled(doc, "2.1 처리방침 주소가 안 열림", 3)
    add_body(
        doc,
        "GitHub Pages로 docs/privacy.html을 올렸다. "
        "저장소가 Private이면 공개 URL이 안 된다. "
        "Settings → Danger zone → Change repository visibility → Public. "
        "Pages 소스는 main, 폴더 / (root). /docs 로 두면 경로가 달라진다.",
    )
    add_body(
        doc,
        "열린 주소: https://dongminhwan.github.io/SubscriptionTracker/docs/privacy.html",
        before=6,
    )

    add_heading_styled(doc, "2.2 Privacy Policy URL 칸을 못 찾음", 3)
    add_body(
        doc,
        "App Information에서 인앱결제 Set Up URL을 열었다. "
        "그건 서버 알림용이라 이 앱과 상관없다. 건드리면 안 된다.",
    )
    add_body(
        doc,
        "처리방침 주소는 버전 페이지의 Privacy Policy URL, "
        "또는 App Privacy의 Edit이다. "
        "처음에는 example.com 자리표시가 들어 있었다. 지우고 위 주소를 넣었다. "
        "User Privacy Choices URL은 비워 두었다.",
        before=6,
    )

    add_heading_styled(doc, "2.3 Save가 회색", 3)
    add_body(
        doc,
        "스크린샷만 넣고 Save를 누르려 하니 버튼이 비활성이었다. "
        "오류 상자 없이, 필수 칸이 비어 있으면 저장이 안 된다. "
        "Description, Keywords, Support URL을 채운 뒤에야 Save가 켜졌다.",
    )

    add_heading_styled(doc, "2.4 스크린샷 크기가 안 맞음", 3)
    add_body(
        doc,
        "6.5인치 칸은 1284×2778을 받는다. "
        "기기 원본 1320×2868은 거절된다. "
        "6.5_ 가 붙은 파일을 올렸다.",
    )

    add_heading_styled(doc, "2.5 Sign-in required가 켜져 있었음", 3)
    add_body(
        doc,
        "체크를 켜 두면 심사자가 로그인 계정을 찾는다. "
        "이 앱은 계정이 없다. 끄지 않고 내면 반려 사유가 된다. "
        "오류 목록에는 안 나와도, 켜져 있는 것 자체가 실수다.",
    )

    add_heading_styled(doc, "2.6 App Privacy를 발행하지 않음", 3)
    add_body(
        doc,
        "처리방침 URL만 있으면 되는 줄 알았다. "
        "Get Started 설문이 비어 있으면 제출이 막히거나 거절된다. "
        "이 앱은 Data Not Collected. 답한 뒤 Publish를 눌러야 한다.",
    )

    add_heading_styled(doc, "2.7 Copyright 칸을 다른 데 찾음", 3)
    add_body(
        doc,
        "연령 등급이 끝난 뒤 App Information의 URL 칸에 "
        "저작권을 넣으려 했다. 그 칸이 아니다. "
        "Copyright는 1.0 Prepare for Submission 페이지에 있다.",
    )

    add_heading_styled(doc, "2.8 Free가 목록에 없음", 3)
    add_body(
        doc,
        "App Pricing에서 Price가 Choose로 되어 있고 Free 글자가 없었다. "
        "무료는 목록 맨 위 0 / ₩0.00 이다. "
        "기준 나라를 한국으로 두면 ₩0.00, 미국이면 $0.00. 같은 뜻이다.",
    )

    add_heading_styled(doc, "3. 오류처럼 보였지만 오류가 아닌 것", 2)
    make_table(
        doc,
        ["보인 것", "실제"],
        [
            [
                "스크린샷 올린 뒤, 다른 아이폰 크기에도 이 사진을 쓰겠느냐는 창",
                "안내다. 오류가 아니다. 이 앱은 아이폰만이라 OK",
            ],
            [
                "Save가 회색 (가격·가용성을 막 끝낸 뒤)",
                "이미 저장된 상태. 다시 누를 변경이 없으면 회색이 정상",
            ],
            [
                "나라마다 Available on App Release, 노란 시계, 24시간 안내",
                "출시 전이므로 “출시되면 이 나라에 뜹니다”라는 뜻. 기다릴 필요 없음",
            ],
            [
                "Draft Submission의 빨간 빼기 버튼",
                "제출 항목을 빼는 버튼. 오류가 아니다. 누르면 안 됨",
            ],
            [
                "1 Item Submitted, 최대 48시간",
                "접수가 된 것이다. 실패가 아님",
            ],
        ],
        [7.8, 8.8],
    )

    add_heading_styled(doc, "4. 빌드/아카이브에서 난 것", 2)
    add_body(
        doc,
        "Connect 빨간 목록은 아니지만, 오늘 위젯 빌드 2를 올리다가 막혔다.",
    )
    add_number(
        doc,
        "스킴이 SubscriptionWidgetExtension으로 잡혀 있었다. "
        "올려야 하는 것은 앱 본체 SubscriptionTracker다.",
    )
    add_number(
        doc,
        "Xcode가 Multiple commands produce / 패키지 꼬임처럼 잠시 실패했다. "
        "앱 코드 문제가 아니라, 방금 빌드와 아카이브가 겹치면 난다. "
        "Xcode를 완전히 종료(⌘Q)한 뒤 스킴을 바꾸고 다시 Archive 했다.",
    )
    add_number(
        doc,
        "빌드 1은 위젯 없음, 빌드 2가 위젯 있음. "
        "심사에 고른 것은 1.0.0 (2)다.",
    )

    add_heading_styled(doc, "5. 앱을 처음 등록할 때 (8월 19일)", 2)
    add_body(
        doc,
        "오늘 제출 오류는 아니지만, 스토어 자리를 만들 때 한 번 막혔다. "
        "같은 종류의 “칸이 거절된” 경험이라 적어 둔다.",
    )
    add_body(
        doc,
        "Name에 구독 트래커를 넣으니 "
        "The app name you entered is already being used. "
        "홈 화면 이름은 그대로 두고, 스토어 이름만 내 구독 트래커로 바꿨다.",
        before=6,
    )
    add_body(
        doc,
        "번들 ID 드롭다운에 위젯 ID가 같이 나온다. "
        "앱 자리를 만들 때는 com.dm.SubscriptionTracker 만 고른다.",
        before=6,
    )

    add_heading_styled(doc, "6. 다음에 같은 빨간 목록이 나오면", 2)
    add_number(doc, "버전 페이지의 연락처 네 칸부터 채우고 Save")
    add_number(doc, "Unable to Add for Review에 적힌 항목만 사이드바에서 연다")
    add_number(
        doc,
        "카테고리·권리·연령은 App Information, "
        "Copyright는 버전 페이지, 가격은 Pricing",
    )
    add_number(doc, "한 칸 고치고 전체를 다시 읽지 않는다. 목록이 줄어드는지 보면 된다")
    add_number(doc, "Free, 24시간, 노란 시계는 이 문서 3장을 먼저 본다")

    os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
    doc.save(OUT_PATH)


if __name__ == "__main__":
    build()
    print(OUT_PATH)
