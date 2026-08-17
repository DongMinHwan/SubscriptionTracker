#!/usr/bin/env python3
"""구독트래커 v1.0.1 추가 기획서(런치스크린 / 앱 아이콘) 생성."""

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

from build_docx import (
    add_body,
    add_heading_styled,
    configure_document,
    make_table,
    set_paragraph_format,
    style_run,
)

OUT_DIR = "/Users/hwangdongmin/Projects/NewProject_1/기획서/v1.0.1"
OUT = os.path.join(OUT_DIR, "구독트래커_v1.0.1_기획서.docx")

ASSETS = "/Users/hwangdongmin/.cursor/projects/Users-hwangdongmin-Projects-NewProject-1/assets"
APP_ICON = os.path.join(ASSETS, "app-icon.png")
LAUNCH_LOGO = os.path.join(ASSETS, "launch-logo.png")

KV = [4.0, 12.6]


def add_image_pair(doc):
    table = doc.add_table(rows=2, cols=2)
    table.autofit = False

    captions = ["앱 아이콘 (배경 #1F4E79, 흰 마크)", "런치스크린 (배경 #F2F2F7, 네이비 마크)"]
    images = [APP_ICON, LAUNCH_LOGO]

    for index, (image, caption) in enumerate(zip(images, captions)):
        cell = table.rows[0].cells[index]
        cell.width = Cm(8.3)
        paragraph = cell.paragraphs[0]
        set_paragraph_format(paragraph, before=0, after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
        paragraph.add_run().add_picture(image, width=Cm(5.0))

        caption_cell = table.rows[1].cells[index]
        caption_cell.width = Cm(8.3)
        caption_paragraph = caption_cell.paragraphs[0]
        set_paragraph_format(caption_paragraph, before=0, after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        style_run(caption_paragraph.add_run(caption), 11, False)


def build():
    os.makedirs(OUT_DIR, exist_ok=True)

    doc = Document()
    configure_document(doc)

    add_heading_styled(doc, "구독 트래커 v1.0.1 기획서 — 런치스크린 / 앱 아이콘", 1)

    make_table(
        doc,
        ["항목", "내용"],
        [
            ["버전", "v1.0.1"],
            ["기준 문서", "구독 트래커 v1.0.0 기획서"],
            ["작성일", "2026년 8월 17일"],
            ["추가하는 것", "런치스크린 1장, 앱 아이콘 1개, 로고 마크 1개"],
            ["v1.0.0 변경", "없음. 화면 3개(목록·추가·상세), 필드 4개, 데이터 규칙은 그대로."],
            ["구현 원칙", "이 문서와 v1.0.0에 적힌 것만 구현한다."],
        ],
        KV,
    )

    add_heading_styled(doc, "1. 왜 추가하는가", 2)
    add_body(
        doc,
        "v1.0.0은 런치스크린을 정하지 않았다. 그래서 Xcode가 자동 생성한 빈 런치스크린이 뜨고, 앱을 켜면 흰 화면이 잠깐 보인다.",
    )
    add_body(
        doc,
        "앱 아이콘도 비어 있어 홈 화면에 회색 사각형으로 보인다.",
    )
    add_body(
        doc,
        "이 두 가지만 채운다. 화면과 기능은 늘리지 않는다.",
    )

    add_heading_styled(doc, "2. 범위 / 비범위", 2)
    make_table(
        doc,
        ["구분", "항목", "내용"],
        [
            ["범위", "런치스크린", "배경색 + 가운데 로고 마크. 정적 1장."],
            ["범위", "앱 아이콘", "1024×1024 한 장."],
            ["범위", "로고 마크", "런치스크린과 앱 아이콘에 같이 쓰는 그림 하나."],
            ["범위", "디자인 토큰", "런치·아이콘용 값 추가. v1.0.0 토큰은 바꾸지 않는다."],
            ["비범위", "스플래시", "페이드, 애니메이션, 일부러 지연시키는 시작 화면 없음."],
            ["비범위", "온보딩", "v1.0.0과 동일하게 없음."],
            ["비범위", "글자", "런치스크린에 앱 이름, 버전, 로딩 문구를 넣지 않는다."],
            ["비범위", "변형", "다크 모드 아이콘, 틴트 아이콘, 대체 아이콘 없음."],
            ["비범위", "기능", "알림, 위젯, 화면 추가, 필드 추가 없음."],
        ],
        [2.4, 4.2, 10.0],
    )

    add_heading_styled(doc, "3. 로고 마크", 2)
    add_body(
        doc,
        "순환 화살표 안에 원화 기호를 넣은 그림 하나를 쓴다. 매달 반복해서 나가는 돈이라는 뜻이다.",
    )
    add_image_pair(doc)
    make_table(
        doc,
        ["항목", "내용"],
        [
            ["구성", "원형 화살표(반복) + 원화 기호(돈). 두 요소만."],
            ["색", "밝은 배경에서는 #1F4E79 마크. 어두운 배경에서는 흰색 마크."],
            ["쓰는 곳", "런치스크린 가운데, 앱 아이콘. 두 곳뿐이다."],
            ["앱 화면", "목록·추가·상세 화면에는 로고를 넣지 않는다."],
            ["형태", "단색 평면. 그라데이션, 그림자, 입체 효과 없음."],
        ],
        KV,
    )

    add_heading_styled(doc, "4. 런치스크린", 2)
    make_table(
        doc,
        ["항목", "내용"],
        [
            ["방식", "Info.plist의 UILaunchScreen 딕셔너리. 스토리보드 파일을 만들지 않는다."],
            ["배경", "#F2F2F7. v1.0.0의 color.background와 같은 값."],
            ["가운데", "로고 마크 120pt 정사각. 화면 정중앙."],
            ["그 외", "없음. 텍스트, 로딩 인디케이터, 버전 표기 없음."],
            ["다크 모드", "없음. 앱과 같이 라이트 고정."],
            [
                "제약",
                "런치스크린은 앱 코드가 실행되기 전에 시스템이 그리는 정적 화면이다. 애니메이션, 이번 달 총액 같은 데이터, 조건 분기가 불가능하다.",
            ],
            [
                "이유",
                "첫 화면(목록)이 #F2F2F7 배경이라 같은 색을 써서 이어지게 한다. 광고판처럼 만들지 않는다.",
            ],
        ],
        KV,
    )

    add_heading_styled(doc, "5. 앱 아이콘", 2)
    make_table(
        doc,
        ["항목", "내용"],
        [
            ["크기", "1024×1024 PNG 한 장."],
            ["배경", "#1F4E79로 정사각형을 꽉 채운다. 여백 없음."],
            ["마크", "흰색 로고 마크. 가운데."],
            ["투명도", "없음. 알파 채널을 넣지 않는다."],
            ["모서리", "직각으로 만든다. 둥근 모서리는 iOS가 알아서 씌운다."],
            ["변형", "다크·틴트·대체 아이콘 없음."],
        ],
        KV,
    )

    add_heading_styled(doc, "6. 에셋", 2)
    make_table(
        doc,
        ["에셋 이름", "종류", "값"],
        [
            ["AppIcon", "App Icon", "1024×1024, 알파 없음"],
            ["LaunchLogo", "Image Set", "120pt 정사각. @3x 기준 360×360px"],
            ["LaunchBackground", "Color Set", "#F2F2F7"],
        ],
        [4.6, 3.6, 8.4],
    )
    add_body(
        doc,
        "Info.plist의 UILaunchScreen에는 UIColorName으로 LaunchBackground, UIImageName으로 LaunchLogo를 지정한다.",
        before=8,
    )

    add_heading_styled(doc, "7. 디자인 토큰 추가", 2)
    add_body(doc, "v1.0.0 11장 토큰에 아래만 더한다. 기존 값은 바꾸지 않는다.")
    make_table(
        doc,
        ["토큰", "값"],
        [
            ["launch.background", "#F2F2F7"],
            ["launch.logoSize", "120pt"],
            ["icon.background", "#1F4E79"],
            ["icon.mark", "#FFFFFF"],
        ],
        [5.6, 11.0],
    )

    add_heading_styled(doc, "8. 완료 기준", 2)
    make_table(
        doc,
        ["번호", "기준"],
        [
            ["1", "앱을 켜면 흰 화면 대신 #F2F2F7 배경에 로고 마크가 보인다."],
            ["2", "홈 화면에 회색 사각형이 아니라 네이비 배경의 앱 아이콘이 보인다."],
            ["3", "목록·추가·상세 화면은 v1.0.0과 똑같다."],
        ],
        [2.0, 14.6],
    )

    add_body(doc, "이 기획서에 없는 기능은 구현하지 말 것.", before=16)

    doc.save(OUT)


if __name__ == "__main__":
    build()
    print(OUT)
