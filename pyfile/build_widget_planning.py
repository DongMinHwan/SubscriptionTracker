#!/usr/bin/env python3
"""구독트래커 홈 화면 위젯 기획서 생성.

v1.1.0으로 만들다가 “이번 달”의 정의가 정해지지 않아 보류한 문서다.
정의가 서면 이 스크립트의 경로만 되돌려 다음 버전 번호로 다시 낸다.
"""

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

from build_docx import (
    add_body,
    add_bullet,
    add_heading_styled,
    add_number,
    configure_document,
    make_table,
    set_paragraph_format,
    style_run,
)

OUT_DIR = "/Users/hwangdongmin/Projects/NewProject_1/기획서/보류/위젯"
OUT = os.path.join(OUT_DIR, "구독트래커_v1.1.0_기획서.docx")
IMAGES = os.path.join(OUT_DIR, "이미지")

KV = [4.0, 12.6]


def add_caption(doc, text):
    p = doc.add_paragraph()
    set_paragraph_format(p, before=2, after=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    style_run(p.add_run(text), 10, False)


def add_image(doc, filename, caption, width_cm):
    p = doc.add_paragraph()
    set_paragraph_format(p, before=6, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    p.add_run().add_picture(os.path.join(IMAGES, filename), width=Cm(width_cm))
    add_caption(doc, caption)


def build():
    os.makedirs(OUT_DIR, exist_ok=True)

    doc = Document()
    configure_document(doc)

    add_heading_styled(doc, "구독 트래커 v1.1.0 기획서 — 홈 화면 위젯", 1)

    make_table(
        doc,
        ["항목", "내용"],
        [
            ["버전", "v1.1.0"],
            ["기준 문서", "구독 트래커 v1.0.0 기획서, v1.0.1 기획서"],
            ["작성일", "2026년 8월 18일"],
            ["추가하는 것", "홈 화면 위젯 2종(작은 것, 중간 것)"],
            [
                "바뀌는 것",
                "데이터 저장 위치가 앱 전용 공간에서 App Group 공용 공간으로 이동. "
                "목록 화면의 날짜 표기가 “8월 28일”에서 “D-9”로 바뀐다(4장). "
                "화면 구성과 계산 규칙은 그대로.",
            ],
            ["구현 원칙", "이 문서에 적힌 것만 구현한다. 적히지 않은 것은 만들지 않는다."],
        ],
        KV,
    )

    add_heading_styled(doc, "1. 왜 만드는가", 2)
    add_body(
        doc,
        "이유는 두 가지고, 둘 다 v1.0.1을 마친 뒤에 분명해졌다.",
    )

    add_heading_styled(doc, "1.1 기획 원안이 아직 절반이다", 3)
    add_body(
        doc,
        "v1.0.0 기획서의 한 문장 정의는 “매달 나가는 구독을 한곳에 모아, 이번 달 얼마인지 홈 화면에서 바로 보게 한다”였다. "
        "지금은 “한곳에 모으기”만 됐고 “홈 화면에서 바로 보기”가 없다. 앱을 열어야만 총액을 볼 수 있다.",
    )
    add_body(
        doc,
        "흔적도 남아 있다. DesignTokens.swift에는 아직 아무도 쓰지 않는 widgetLabel, widgetValue 토큰이 있다. "
        "처음부터 위젯을 염두에 뒀다가 범위를 좁히며 빠진 자리다.",
    )

    add_heading_styled(doc, "1.2 심사 가이드라인 4.2 대비", 3)
    add_body(
        doc,
        "App Review Guideline 4.2(Minimum Functionality)는 “앱이 특별히 유용하거나, 고유하거나, 앱다운 것이 아니라면 "
        "App Store에 있을 자리가 없다”고 말한다. 반려 사유의 큰 비중이 이 조항에 몰려 있다.",
    )
    add_body(
        doc,
        "현재 앱은 화면 3개짜리 목록·추가·상세와 합계가 전부다. 웹뷰 래퍼가 아니고 오프라인에서 동작하며 데이터가 남는다는 점은 "
        "유리하지만, 심사자가 “이건 메모 앱으로도 되는 것 아닌가”라고 물으면 지금은 답할 근거가 약하다.",
    )
    add_body(
        doc,
        "위젯은 웹사이트가 할 수 없는 일이고, 이 앱의 핵심 사용 흐름에 직접 붙는다. 장식으로 얹는 기능이 아니라 "
        "원래 정의의 나머지 절반이라는 점이 중요하다.",
    )

    add_heading_styled(doc, "2. 범위", 2)
    make_table(
        doc,
        ["하는 것", "하지 않는 것"],
        [
            [
                "홈 화면 위젯 (작은 것)\n"
                "홈 화면 위젯 (중간 것)\n"
                "App Group으로 데이터 공유\n"
                "앱에서 데이터가 바뀌면 위젯 갱신\n"
                "구독이 없을 때의 빈 상태\n"
                "위젯을 누르면 앱 열기",
                "잠금화면 위젯, StandBy\n"
                "큰 위젯\n"
                "위젯 설정(어떤 구독을 볼지 고르기)\n"
                "결제 전 알림\n"
                "차트, 통계\n"
                "Live Activity, Siri, App Intents",
            ],
        ],
        [8.3, 8.3],
    )
    add_body(
        doc,
        "하지 않는 것 목록이 긴 이유는 v1.0.0과 같다. 넣을 수 있는 것과 넣어야 하는 것은 다르다. "
        "알림은 다음 버전에서 다룬다.",
        before=8,
    )

    add_heading_styled(doc, "3. 위젯 사양", 2)

    add_heading_styled(doc, "3.1 표시 내용", 3)
    add_body(doc, "두 위젯 모두 같은 데이터를 쓰고, 자리가 되는 만큼만 보여준다.")
    make_table(
        doc,
        ["항목", "작은 위젯", "중간 위젯"],
        [
            ["이번 달 총액", "표시 (28pt 굵게)", "표시 (22pt 굵게)"],
            [
                "이번 달 구독",
                "대표 1건의 이름과 나머지 개수 (“네이버맴버십 외 2개”)",
                "개수만 (“구독 5개”)",
            ],
            [
                "다가오는 결제 목록",
                "표시하지 않음",
                "최대 3건. 한 줄에 이름 · D-day · 금액",
            ],
        ],
        [4.0, 6.3, 6.3],
    )
    add_body(
        doc,
        "작은 위젯에 결제일을 적지 않는 이유가 있다. 오늘이 결제일인 구독은 “다음 결제 8월 18일”처럼 보이는데, "
        "앱은 오늘 것이 이미 빠져나갔는지 알 수 없다. 아직 나갈 돈인 것처럼 읽히므로 작은 위젯은 금액만 말한다.",
        before=8,
    )
    add_body(
        doc,
        "개수 대신 대표 이름을 부르는 이유도 같은 결이다. “구독 3개”는 무엇을 센 건지 와닿지 않는다. "
        "대표는 결제일이 가장 이른 것을 고르므로, 중간 위젯 목록의 첫 줄과 항상 같은 이름이 나온다.",
    )
    add_body(
        doc,
        "여기서 세는 개수는 이번 달 총액에 들어간 구독만이다. 다른 달에 결제되는 매년 구독은 총액에서 빠지므로 "
        "개수에서도 빠진다. 두 숫자가 다른 기준을 쓰면 “₩40,000 / 구독 4개”처럼 서로 안 맞는 값이 나란히 놓인다.",
    )

    add_heading_styled(doc, "3.2 시안", 3)
    add_body(
        doc,
        "실제 위젯을 만들기 전에 배치를 합의하려고 그린 그림이다. 색과 글자 크기는 아래 토큰 값을 그대로 썼다.",
    )
    add_image(doc, "위젯_작은것.png", "작은 위젯 — 총액과 다음 결제 1건", 5.5)
    add_image(doc, "위젯_중간것.png", "중간 위젯 — 왼쪽 총액, 오른쪽 다음 결제", 12.0)
    add_image(doc, "위젯_빈상태.png", "구독이 하나도 없을 때", 5.5)

    add_heading_styled(doc, "3.3 빈 상태", 3)
    add_body(
        doc,
        "총액이 ₩0인 경우가 두 가지인데, 사용자가 해야 할 일이 서로 달라 문구를 나눈다. "
        "위젯을 비워 두면 고장 난 것처럼 보이므로 어느 쪽이든 무언가는 적는다.",
    )
    make_table(
        doc,
        ["상황", "문구"],
        [
            ["구독이 하나도 없다", "“구독을 추가하면 여기에 보입니다” (작은 위젯은 두 줄로 접는다)"],
            [
                "구독은 있는데 이번 달 결제가 없다",
                "“이번 달 결제 없음”. 다른 달에 결제되는 매년 구독만 있을 때 나온다. "
                "여기에 추가 안내를 띄우면 이미 넣은 사람에게 안 넣었다고 말하는 셈이 된다.",
            ],
        ],
        [4.6, 12.0],
    )

    add_heading_styled(doc, "3.4 누르면 하는 일", 3)
    add_body(
        doc,
        "위젯 어디를 눌러도 앱의 목록 화면이 열린다. 특정 구독의 상세로 바로 가는 딥링크는 이번 범위가 아니다.",
    )

    add_heading_styled(doc, "3.5 언제 갱신하는가", 3)
    make_table(
        doc,
        ["시점", "방법"],
        [
            [
                "앱에서 구독을 추가·수정·삭제했을 때",
                "저장 직후 WidgetCenter로 갱신을 요청한다. 사용자가 홈 화면으로 돌아왔을 때 이미 반영돼 있어야 한다.",
            ],
            [
                "날짜가 바뀔 때",
                "타임라인의 다음 항목을 다음 날 자정으로 잡는다. 날이 바뀌면 “다음 결제”가 달라지고, 달이 바뀌면 총액이 달라진다.",
            ],
            [
                "그 밖",
                "따로 하지 않는다. 위젯 갱신 횟수는 시스템이 제한하므로 자주 요청할수록 오히려 늦어진다.",
            ],
        ],
        [5.0, 11.6],
    )

    add_heading_styled(doc, "4. 날짜 표기 (앱·위젯 공통)", 2)
    add_body(
        doc,
        "v1.0.0에서는 목록 행에 “3월 21일” 형식으로 적었다. v1.1.0부터 앱과 위젯 모두 D-day로 바꾼다.",
    )

    add_heading_styled(doc, "4.1 규칙", 3)
    make_table(
        doc,
        ["상황", "표기", "예"],
        [
            ["결제일이 남았다", "D-남은일수", "D-9"],
            ["결제일이 오늘이다", "D-DAY", "D-DAY"],
        ],
        [5.3, 6.0, 5.3],
    )
    add_body(
        doc,
        "주기에 따라 다르게 적지 않는다. 매달이든 매년이든 형식이 같다.",
        before=8,
    )
    add_body(
        doc,
        "남은 일수는 자정 기준으로 센다. 오늘과 결제일의 날짜 차이이며 시각은 보지 않는다. "
        "결제일이 지나면 다음 주기로 밀리므로 음수가 나올 자리는 없다.",
    )

    add_heading_styled(doc, "4.2 왜 바꾸는가", 3)
    add_body(
        doc,
        "“8월 28일”을 읽으려면 오늘이 며칠인지 떠올려 빼야 한다. 이 앱에서 날짜를 보는 이유는 "
        "“언제 나가는지”지 “며칠인지”가 아니다. 그 뺄셈을 앱이 대신한다.",
    )
    add_body(
        doc,
        "약점도 알고 간다. D-364는 언제인지 바로 그려지지 않는다. 그래도 형식을 하나로 두는 쪽을 골랐다. "
        "어떤 줄은 날짜고 어떤 줄은 D-day면 눈이 두 번 적응해야 한다. 정확한 날짜가 필요하면 행을 눌러 상세에서 본다.",
    )

    add_heading_styled(doc, "4.3 주기는 목록에 적지 않는다", 3)
    add_body(
        doc,
        "한때 매년 구독에만 “매년”을 붙이려 했다. 129,000원 같은 큰 금액이 한 달에 나가는 줄 읽힐까 봐였다. "
        "날짜로 적을 때는 “매년 8월 18일”이 말이 됐지만, D-day로 바꾸니 “매년 D-364”가 되어 어색해졌다.",
    )
    add_body(
        doc,
        "빼도 잃는 것이 없다. 매달 구독은 다음 결제가 늘 한 달 안이라 D-31을 넘을 수 없다. "
        "D-364라는 숫자 자체가 이미 매년 구독이라고 말한다. 정확한 주기는 행을 눌러 상세에서 본다.",
    )

    add_heading_styled(doc, "5. 데이터 공유 (App Group)", 2)

    add_heading_styled(doc, "5.1 왜 필요한가", 3)
    add_body(
        doc,
        "위젯은 앱과 별도의 프로세스로 돌아간다. 지금 SwiftData 저장소는 앱 전용 폴더에 있어서 위젯이 읽을 수 없다. "
        "두 프로세스가 같은 파일을 보게 하려면 App Group이라는 공용 공간으로 저장소를 옮겨야 한다.",
    )

    add_heading_styled(doc, "5.2 식별자", 3)
    make_table(
        doc,
        ["항목", "값"],
        [
            ["앱 번들 ID", "com.dm.SubscriptionTracker (그대로)"],
            ["위젯 번들 ID", "com.dm.SubscriptionTracker.SubscriptionWidget"],
            ["App Group", "group.com.dm.SubscriptionTracker"],
            ["대상", "앱 타깃과 위젯 타깃 양쪽에 같은 App Group을 켠다."],
        ],
        KV,
    )

    add_heading_styled(doc, "5.3 기존 데이터는 초기화된다", 3)
    add_body(
        doc,
        "저장 위치가 바뀌므로 이전 위치에 있던 데이터는 앱에서 보이지 않게 된다. 옮기는 코드는 만들지 않는다.",
    )
    add_body(
        doc,
        "아직 출시 전이라 실제 사용자가 없기 때문에 지금이 바꾸기 가장 싼 시점이다. "
        "출시 후에 바꾸려면 마이그레이션 코드를 따로 써야 한다. 내 시뮬레이터에 넣어 둔 샘플 데이터는 다시 넣으면 된다.",
    )

    add_heading_styled(doc, "6. 계산 규칙", 2)
    add_body(doc, "총액 규칙은 새로 정하지 않는다. 앱에 이미 있는 것을 위젯이 그대로 쓴다.")
    make_table(
        doc,
        ["항목", "규칙"],
        [
            [
                "이번 달 총액",
                "MonthlyTotal.amount를 그대로 쓴다. 매달 구독은 전부 더하고, 매년 구독은 결제월이 이번 달일 때만 더한다.",
            ],
            [
                "이번 달 구독 수",
                "총액에 들어간 구독만 센다. 판정은 총액과 같은 함수(MonthlyTotal.includes)를 쓴다.",
            ],
            [
                "다가오는 결제",
                "정렬은 목록 화면과 같다. 다음 결제일 오름차순, 같은 날이면 이름 가나다순. 중간 위젯은 앞의 3건까지.",
            ],
            [
                "지난 날짜 처리",
                "위젯은 저장을 할 수 없으므로 upcomingPaymentDate로 밀어서 계산만 한다. "
                "앱을 오래 안 열어 저장된 결제일이 과거로 남아 있어도 같은 값이 나온다.",
            ],
            ["통화 표기", "Formatters의 원화 표기를 그대로 쓴다. 소수점 없음, 천 단위 쉼표."],
        ],
        [4.0, 12.6],
    )
    add_body(
        doc,
        "위젯이 데이터를 고치는 일은 없다. 읽기 전용이다.",
        before=8,
    )

    add_heading_styled(doc, "7. 디자인 토큰", 2)
    add_body(
        doc,
        "DesignTokens.swift에 이미 있는 두 개를 드디어 쓰고, 위젯에만 필요한 것을 더한다.",
    )
    make_table(
        doc,
        ["토큰", "값", "쓰는 곳"],
        [
            ["Typography.widgetLabel", "Font.caption (12pt)", "“이번 달”, 구독 수, D-day (이미 있음)"],
            ["Typography.widgetValue", "Font.title2.bold() (22pt)", "중간 위젯의 총액 (이미 있음)"],
            ["Typography.widgetValueLarge", "Font.title.bold() (28pt)", "작은 위젯의 총액 (추가)"],
            ["Typography.widgetName", "Font.subheadline (15pt)", "구독 이름, 금액 (추가)"],
            ["Widget.background", "라이트 #FFFFFF / 다크 #1C1C1E", "컨테이너 배경 (추가)"],
            ["Widget.text", "라이트 #1C1C1E / 다크 #FFFFFF", "이름, 금액 (추가)"],
            ["Widget.textSecondary", "라이트 #8E8E93 / 다크 #98989F", "라벨, D-day (추가)"],
            ["Widget.accent", "라이트 #1F4E79 / 다크 #6FA8DC", "총액 글자색 (추가)"],
            ["Widget.separator", "라이트 #C6C6C8 / 다크 #38383A", "중간 위젯의 구분선 (추가)"],
        ],
        [5.2, 5.4, 6.0],
    )
    add_body(
        doc,
        "총액이 작은 위젯에서 더 큰 이유는 그 위젯의 내용이 총액뿐이기 때문이다. 중간 위젯에서는 총액이 머리글 한 줄에 "
        "들어가고 아래 목록이 주인공이라 한 단계 작게 쓴다.",
        before=8,
    )
    add_body(
        doc,
        "앱 본체는 preferredColorScheme(.light)로 라이트에 고정돼 있지만, 위젯은 홈 화면 배경 위에 놓이므로 "
        "시스템 설정을 따른다. 그래서 위젯 색만 라이트·다크 두 벌을 둔다. "
        "에셋 카탈로그 대신 UIColor의 트레잇 초기화로 만든다. 토큰이 한 파일에 모여 있어야 값을 대조하기 쉽다.",
    )

    add_heading_styled(doc, "8. 작업 순서와 담당", 2)
    add_body(
        doc,
        "위젯 타깃 추가는 Xcode에서 직접 하는 편이 안전하다. 프로젝트 파일을 손으로 고치면 깨지기 쉽다.",
    )
    make_table(
        doc,
        ["순서", "할 일", "누가"],
        [
            [
                "1",
                "Xcode에서 File → New → Target → Widget Extension 추가. 이름 SubscriptionWidget, Live Activity와 Configuration Intent는 끈다.",
                "직접",
            ],
            [
                "2",
                "앱 타깃과 위젯 타깃의 Signing & Capabilities에서 App Groups를 켜고 group.com.dm.SubscriptionTracker를 추가.",
                "직접",
            ],
            ["3", "공유 코드를 Shared 폴더로 옮기고 앱·위젯 두 타깃에 모두 넣는다", "코드"],
            ["4", "ModelConfiguration을 App Group 컨테이너로 옮기고 위젯과 공유", "코드"],
            ["5", "위젯 타임라인과 두 크기의 화면 구현", "코드"],
            ["6", "앱에서 저장할 때 위젯 갱신 요청 추가", "코드"],
            ["7", "앱 목록과 위젯의 날짜 표기를 D-day로 바꾼다 (4장)", "코드"],
            ["8", "시뮬레이터에서 위젯 추가·갱신·빈 상태 확인", "같이"],
        ],
        [1.6, 11.7, 3.3],
    )

    add_heading_styled(doc, "9. 완료 기준", 2)
    add_body(doc, "아래가 전부 되면 v1.1.0을 끝난 것으로 본다.")
    add_number(doc, "홈 화면에 작은 위젯과 중간 위젯을 추가할 수 있다.")
    add_number(doc, "앱에서 구독을 추가하고 홈 화면으로 나오면 위젯의 총액이 바뀌어 있다.")
    add_number(doc, "구독을 전부 지우면 위젯이 빈 상태 문구를 보여준다. 빈 사각형이 아니다.")
    add_number(doc, "위젯을 누르면 앱이 열린다.")
    add_number(doc, "다크 모드에서 총액 글자가 배경에 묻히지 않는다.")
    add_number(doc, "앱을 지웠다 다시 깔아도 위젯이 죽지 않는다.")
    add_number(doc, "앱 목록과 중간 위젯이 같은 D-day 형식으로 날짜를 보여준다.")
    add_number(doc, "다른 달에 결제되는 매년 구독만 남기면 “이번 달 결제 없음”이 나온다.")
    add_number(doc, "iOS 18 시뮬레이터에서 위 항목을 모두 확인했다.")

    add_heading_styled(doc, "10. 심사 제출 시 적을 것", 2)
    add_body(
        doc,
        "App Store Connect의 심사 메모(App Review Notes)는 비워 두지 않는다. 4.2로 걸리는 앱 상당수가 "
        "“무엇이 네이티브인지” 설명하지 않아서 걸린다.",
    )
    add_bullet(doc, "홈 화면 위젯 2종으로 앱을 열지 않고도 이번 달 고정지출과 다음 결제를 확인할 수 있다는 점.")
    add_bullet(doc, "네트워크 없이 전부 동작하고, 데이터가 기기 안에만 남는다는 점.")
    add_bullet(doc, "계정 가입이 없어 심사자가 바로 모든 기능을 쓸 수 있다는 점.")
    add_body(
        doc,
        "이 메모는 v1.1.0 구현이 끝난 뒤 실제 화면을 보고 다시 다듬는다. 없는 기능을 적으면 그게 더 큰 문제가 된다.",
        before=8,
    )

    add_heading_styled(doc, "11. 다음 버전 후보", 2)
    add_body(
        doc,
        "이번에 하지 않기로 한 것 중, 다음에 가장 먼저 볼 것들이다. 여기 적어 두는 이유는 지금 하지 않기 위해서다.",
    )
    add_bullet(doc, "결제 전날 로컬 알림. v1.0.0 기획서의 원래 v1 기능 중 마지막으로 남은 하나.")
    add_bullet(doc, "위젯에서 특정 구독 상세로 바로 가는 딥링크.")
    add_bullet(doc, "잠금화면 위젯.")

    doc.save(OUT)


if __name__ == "__main__":
    build()
    print(OUT)
