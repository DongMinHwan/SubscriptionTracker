#!/usr/bin/env python3
"""2026-08-22 개발일지 .docx 생성."""

import os
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from build_docx import (
    add_body,
    add_bullet,
    add_heading_styled,
    add_number,
    configure_document,
    make_table,
)
from docx import Document

BASE = "/Users/hwangdongmin/Projects/NewProject_1"
OUT_PATH = f"{BASE}/개발일지/2026-08-22_개발일지.docx"
KV = [3.2, 13.4]


def build():
    doc = Document()
    configure_document(doc)

    add_heading_styled(doc, "개발일지 — 2026년 8월 22일 (토)", 1)

    make_table(
        doc,
        ["항목", "내용"],
        [
            ["날짜", "2026년 8월 22일 오전~낮"],
            ["프로젝트", "구독 트래커 / 스토어 이름 내 구독 트래커"],
            [
                "목표",
                "처리방침을 만들고, 위젯이 들어간 빌드를 올린 뒤 "
                "App Store 심사를 제출한다.",
            ],
            [
                "결과",
                "1.0.0 (2)를 12:05에 Submit for Review 했다. "
                "출시는 수동이라 승인되어도 스토어에 바로 올라가지 않는다.",
            ],
            [
                "한 줄 소감",
                "앱이 없어서 막힌 게 아니라, Connect 칸이 비어 있어서 막혔다. "
                "빨간 목록은 한 번에 고치지 말고 사이드바 항목을 하나씩 열면 된다.",
            ],
        ],
        KV,
    )

    add_heading_styled(doc, "1. 오늘의 흐름", 2)
    make_table(
        doc,
        ["시각", "한 일", "결과"],
        [
            [
                "10:00 ~ 10:20",
                "개인정보 처리방침 페이지, GitHub Pages",
                "저장소를 Public으로 바꾼 뒤에야 주소가 열렸다",
            ],
            [
                "10:20 ~ 11:30",
                "위젯 임베드 복구, 헤더를 월 구독 금액으로, 아카이브",
                "1.0.0 (2) TestFlight. 스킴이 위젯으로 잡혀 한 번 실패",
            ],
            [
                "11:30 ~ 11:42",
                "스크린샷·설명·키워드·심사 메모·수동 출시",
                "6.5인치 세 장, Sign-in required 끔",
            ],
            [
                "11:42 ~ 11:45",
                "App Privacy 설문 후 첫 Add for Review",
                "카테고리·권리·연령·저작권·가격·연락처가 비어 제출 거부",
            ],
            [
                "11:45 ~ 12:05",
                "빠진 칸을 하나씩 채움",
                "Finance, 4+, ₩0, 연락처. 다시 제출",
            ],
            [
                "12:05 ~ 12:06",
                "Submit for Review",
                "1 Item Submitted. 심사 대기",
            ],
        ],
        [3.4, 7.6, 5.6],
    )
    add_body(
        doc,
        "시각은 대화와 스크린샷 기준이다. "
        "제출 절차와 그때 난 오류는 개발일지에 길게 쓰지 않았다. "
        "출시 폴더에 문서를 따로 두었다.",
        before=8,
    )

    add_heading_styled(doc, "2. 따로 둔 문서", 2)
    make_table(
        doc,
        ["파일", "내용"],
        [
            [
                "출시/2026-08-22_심사제출과정.docx",
                "다음에 같은 앱을 낼 때 누를 순서와 실제로 넣은 값",
            ],
            [
                "출시/2026-08-22_심사제출_오류.docx",
                "첫 Add for Review에서 난 빨간 목록과, 오류처럼 보인 것들",
            ],
        ],
        [7.4, 9.2],
    )

    add_heading_styled(doc, "3. 개인정보 처리방침", 2)
    add_body(
        doc,
        "Connect는 처리방침을 앱 안 PDF가 아니라 https 주소로 받는다. "
        "이 앱은 서버가 없어도 페이지는 있어야 한다. "
        "계정 없음, 기기에만 저장, 알림은 기기 안, 문의 메일.",
    )
    add_body(
        doc,
        "페이지는 docs/privacy.html 이다. "
        "GitHub Pages는 저장소가 Private이면 안 열린다. "
        "오늘 SubscriptionTracker를 Public으로 바꿨다. "
        "Pages 소스는 main, 폴더 / (root).",
        before=6,
    )
    add_body(
        doc,
        "주소: https://dongminhwan.github.io/SubscriptionTracker/docs/privacy.html",
        before=6,
    )
    add_body(
        doc,
        "문의 메일은 hdmin8712@gmail.com 이다. "
        "TestFlight Apple ID(sizzle@naver.com)와 스토어 연락처를 나눴다.",
        before=6,
    )

    add_heading_styled(doc, "4. 위젯을 다시 넣음", 2)
    add_body(
        doc,
        "8월 18일에 “이번 달” 정의가 흔들려 위젯 임베드를 빼 두었다. "
        "8월 19일에 총액 규칙을 정했다. "
        "다음 결제일의 연·월이 이번 달과 같을 때만 더한다. "
        "라벨은 8월 구독 금액.",
    )
    add_body(
        doc,
        "오늘 project.pbxproj에 SubscriptionWidgetExtension.appex 임베드를 다시 넣었다. "
        "위젯 헤더도 DateFormat.monthlyTotalLabel을 쓰게 맞췄다. "
        "갤러리 이름은 구독 금액.",
        before=6,
    )
    add_body(
        doc,
        "아카이브할 때 스킴이 SubscriptionWidgetExtension으로 잡혀 있었다. "
        "Multiple commands produce 비슷한 꼬임이 한 번 났다. "
        "Xcode를 끄고 스킴을 SubscriptionTracker로 바꾼 뒤 다시 Archive 했다. "
        "올라간 빌드는 1.0.0 (2). MARKETING_VERSION은 1.0.0 그대로다.",
        before=6,
    )

    add_heading_styled(doc, "5. 스토어 칸을 채움", 2)
    add_body(
        doc,
        "스크린샷은 6.5인치 칸에 1284×2778 세 장이다. "
        "목록, 추가, 설정. 원본 1320×2868은 이 칸에서 거절된다. "
        "다른 아이폰 크기에 그대로 쓰겠다는 안내가 떴고, OK로 넘겼다.",
    )
    add_body(
        doc,
        "설명·키워드·지원 URL을 넣기 전에는 Save가 회색이었다. "
        "Sign-in required는 켜져 있어서 껐다. 로그인 없는 앱인데 켜 두면 심사자가 계정을 찾는다. "
        "출시는 Manually release this version.",
        before=6,
    )
    add_body(
        doc,
        "App Privacy는 “데이터를 수집하지 않음”으로 답하고 Publish까지 눌러야 한다. "
        "설문만 하고 발행을 안 하면 제출이 막힌다.",
        before=6,
    )

    add_heading_styled(doc, "6. 한 번에 안 올라간 이유", 2)
    add_body(
        doc,
        "버전 페이지를 채운 뒤 Add for Review를 눌렀다. "
        "Unable to Add for Review가 나왔다. "
        "카테고리, 콘텐츠 권리, 연령 등급, 저작권, 가격 티어, "
        "심사 연락처(이름·성·메일·전화)가 비어 있었다.",
    )
    add_body(
        doc,
        "앱 정보가 아니라 Connect 앱 자리의 공통 칸이다. "
        "왼쪽 App Information과 Pricing을 열고 하나씩 넣었다. "
        "자세한 문구와 고친 방법은 오류 문서에 있다.",
        before=6,
    )

    add_heading_styled(doc, "7. 오늘 배운 것", 2)
    add_number(
        doc,
        "처리방침 URL은 앱 기능이 아니라 제출 열쇠다. "
        "GitHub Pages는 저장소가 공개여야 한다.",
    )
    add_number(
        doc,
        "Privacy Policy URL, Copyright는 App Information이 아니라 "
        "버전 페이지에 있다. 인앱결제 Set Up URL과는 다른 칸이다.",
    )
    add_number(
        doc,
        "무료는 Free라는 글자가 아니라 가격 목록의 ₩0.00 이다. "
        "기준 나라는 한국(KRW)로 두었다.",
    )
    add_number(
        doc,
        "Available on App Release의 노란 시계는 출시 전 정상이다. "
        "24시간은 이미 올라간 앱을 바꿀 때 이야기다.",
    )
    add_number(
        doc,
        "아카이브 스킴은 앱 본체여야 한다. 위젯 스킴으로 올리면 안 된다.",
    )
    add_number(
        doc,
        "첫 출시는 수동 출시가 맞다. 1 Item Submitted는 심사가 시작된 것이지 "
        "스토어에 올라간 것이 아니다.",
    )

    add_heading_styled(doc, "8. 남은 것 / 다음에 할 일", 2)
    make_table(
        doc,
        ["순서", "할 일", "메모"],
        [
            [
                "1",
                "심사 결과를 기다린다",
                "메일로 온다. 최대 안내 48시간. 심사 중에는 빌드·설명을 건드리지 않는다",
            ],
            [
                "2",
                "승인되면 Release This Version을 누른다",
                "수동 출시. 안 누르면 스토어에 안 나온다",
            ],
            [
                "3",
                "거절되면 사유를 보고 고친다",
                "오류 문서에 오늘 막힌 칸은 이미 적혀 있다",
            ],
            [
                "4",
                "위젯·스크린샷 코드를 저장한다",
                "임베드 복구와 ScreenshotLaunch는 아직 커밋이 안 됐을 수 있다",
            ],
            [
                "5",
                "결제 이력 → 통계",
                "월말에 총액이 ₩0이 되는 한계. 다음 버전",
            ],
        ],
        [1.6, 6.6, 8.4],
    )

    add_heading_styled(doc, "9. 오늘 만진 파일", 2)
    make_table(
        doc,
        ["파일", "무엇을"],
        [
            ["docs/privacy.html", "처리방침. 문의 hdmin8712@gmail.com"],
            ["project.pbxproj", "위젯 확장 임베드를 다시 넣음"],
            ["SubscriptionWidget.swift", "헤더를 monthlyTotalLabel로, 갤러리 이름"],
            ["ScreenshotLaunch.swift", "스토어 스크린샷용 시드 (DEBUG)"],
            ["기획서/v1.1.0/스크린샷/", "6.5인치 세 장"],
            ["출시/*.docx", "제출 과정과 오류를 일지와 분리"],
        ],
        [7.2, 9.4],
    )

    os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
    doc.save(OUT_PATH)


if __name__ == "__main__":
    build()
    print(OUT_PATH)
