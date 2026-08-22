#!/usr/bin/env python3
"""2026-08-22 심사 제출 과정 .docx 생성."""

import os
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from build_docx import (
    add_body,
    add_bullet,
    add_heading_styled,
    add_number,
    configure_document,
    make_table,
)
from docx import Document

BASE = "/Users/hwangdongmin/Projects/NewProject_1"
OUT_PATH = f"{BASE}/출시/2026-08-22_심사제출과정.docx"
KV = [3.6, 13.0]


def build():
    doc = Document()
    configure_document(doc)

    add_heading_styled(doc, "내 구독 트래커 — 심사 제출 과정", 1)

    make_table(
        doc,
        ["항목", "내용"],
        [
            ["작성일", "2026년 8월 22일"],
            ["앱", "내 구독 트래커 (홈 화면 이름: 구독 트래커)"],
            ["Apple ID", "6803106914"],
            ["번들 ID", "com.dm.SubscriptionTracker"],
            ["제출 빌드", "1.0.0 (2). 위젯 포함"],
            ["제출 시각", "2026년 8월 22일 12:05"],
            ["출시 방식", "Manually release this version"],
            [
                "이 문서",
                "다음에 같은 칸을 채울 때의 순서와, 이번에 실제로 넣은 값. "
                "첫 제출에서 난 빨간 오류는 별도 문서.",
            ],
        ],
        KV,
    )

    add_heading_styled(doc, "1. 제출과 출시는 다른 버튼이다", 2)
    add_body(
        doc,
        "Add for Review / Submit for Review는 Apple 심사에 넣는 것이다. "
        "1 Item Submitted가 나와도 스토어에 앱이 있는 것이 아니다.",
    )
    add_body(
        doc,
        "이번 버전은 수동 출시다. 승인 메일이 오면 "
        "Release This Version을 직접 눌러야 스토어에 보인다. "
        "자동 출시로 두면 통과 즉시 공개된다.",
        before=6,
    )

    add_heading_styled(doc, "2. 제출 전에 채울 순서", 2)
    add_body(
        doc,
        "버전 페이지만 채우고 Add for Review를 누르면 막힌다. "
        "앱 공통 칸이 왼쪽 사이드바에 흩어져 있다. 이 순서가 맞다.",
    )
    make_table(
        doc,
        ["순서", "어디", "무엇을"],
        [
            [
                "1",
                "App Information",
                "Primary Category, Content Rights, Age Ratings",
            ],
            [
                "2",
                "1.0 Prepare for Submission",
                "스크린샷, 설명, 키워드, Support URL, Privacy Policy URL, "
                "Copyright, 빌드, 심사 메모, 연락처, 로그인 체크, 출시 방식",
            ],
            [
                "3",
                "App Privacy",
                "수집 여부 설문 후 Publish",
            ],
            [
                "4",
                "Pricing and Availability",
                "가격 ₩0, 판매 국가",
            ],
            [
                "5",
                "버전 페이지",
                "Save 후 Add for Review → Submit for Review",
            ],
        ],
        [1.8, 5.4, 9.4],
    )

    add_heading_styled(doc, "3. 건드리지 않는 칸", 2)
    add_bullet(doc, "App Clip")
    add_bullet(doc, "인앱결제 Server Notification URL / Set Up URL")
    add_bullet(doc, "User Privacy Choices URL")
    add_bullet(doc, "Marketing URL, Promotional Text (없어도 제출됨)")
    add_bullet(doc, "Age Suitability URL, Made for Kids")
    add_bullet(doc, "External Testing (베타 심사가 붙음). 내부 TestFlight만")

    add_heading_styled(doc, "4. App Information", 2)

    add_heading_styled(doc, "4.1 Category", 3)
    add_body(
        doc,
        "General 아래 App Information. 버전 페이지가 아니다. "
        "Primary는 Finance. Secondary는 비움.",
    )

    add_heading_styled(doc, "4.2 Content Rights", 3)
    add_body(
        doc,
        "Category 근처 Edit. "
        "이 앱은 남의 영상·음악·기사를 넣지 않는다. "
        "No third-party content / 타사 콘텐츠 없음.",
    )

    add_heading_styled(doc, "4.3 Age Ratings", 3)
    add_body(
        doc,
        "같은 페이지에서 Set Up. 폭력·의료·성·도박 등은 전부 NONE / NO. "
        "Parental Controls, Age Assurance, Unrestricted Web Access, "
        "User-Generated Content, Social Media, Messaging, Advertising 도 NO. "
        "구독 이름은 본인만 보는 메모라 User-Generated Content에 넣지 않았다. "
        "결과 등급 4+. Made for Kids는 고르지 않음. URL 칸은 비움.",
    )

    add_heading_styled(doc, "5. 버전 페이지 (1.0 Prepare for Submission)", 2)

    add_heading_styled(doc, "5.1 스크린샷", 3)
    add_body(
        doc,
        "iPhone 6.5인치 칸. 1284×2778. 세 장, 목록 → 추가 → 설정 순. "
        "파일은 기획서/v1.1.0/스크린샷/ 의 6.5_1_목록.png, 6.5_2_추가.png, 6.5_3_설정.png. "
        "1320×2868 원본은 이 칸에서 거절된다. App Preview 영상은 넣지 않음.",
    )
    add_body(
        doc,
        "올린 뒤 “다른 크기에도 이 사진을 쓰겠느냐”는 안내가 뜬다. "
        "아이폰만이고 화면이 기종마다 다르지 않아 OK.",
        before=6,
    )

    add_heading_styled(doc, "5.2 설명과 검색", 3)
    add_body(doc, "언어 Korean. Promotional Text는 비움.")
    add_body(
        doc,
        "Description",
        before=8,
    )
    add_body(
        doc,
        "매달 빠져나가는 구독을 한곳에 모아 둡니다.\n\n"
        "이번 달에 나갈 금액이 목록 위에 보이고, 각 구독은 결제까지 며칠 남았는지 "
        "D-day로 표시합니다. 결제 하루 전 알림을 켤 수 있고, "
        "홈 화면 위젯으로 금액을 바로 볼 수 있습니다.\n\n"
        "계정이 필요 없습니다. 정보는 이 기기에만 남습니다.",
        before=4,
    )
    add_body(doc, "Keywords (쉼표 뒤 공백 없이, 100자 안)", before=8)
    add_body(
        doc,
        "구독,구독관리,고정지출,정기결제,넷플릭스,알림,위젯,가계부,트래커",
        before=4,
    )
    add_body(doc, "What's New in This Version", before=8)
    add_body(doc, "첫 출시입니다.", before=4)

    add_heading_styled(doc, "5.3 URL", 3)
    make_table(
        doc,
        ["칸", "값"],
        [
            [
                "Privacy Policy URL",
                "https://dongminhwan.github.io/SubscriptionTracker/docs/privacy.html",
            ],
            [
                "Support URL",
                "같은 처리방침 주소. 문의 메일이 그 페이지에 있다",
            ],
            ["User Privacy Choices URL", "비움"],
            ["Marketing URL", "비움"],
        ],
        [5.0, 11.6],
    )
    add_body(
        doc,
        "Privacy Policy URL은 App Privacy의 Edit에도 있다. "
        "App Information의 인앱결제 Set Up URL이 아니다.",
        before=8,
    )

    add_heading_styled(doc, "5.4 Copyright", 3)
    add_body(
        doc,
        "버전 페이지에 있다. App Information의 연령 URL 칸이 아니다. "
        "연도만 적으면 ©는 애플이 붙인다. 스토어에 공개되는 한 줄이라 실명은 안 넣었다.",
    )
    add_body(doc, "2026 내 구독 트래커", before=4)

    add_heading_styled(doc, "5.5 빌드", 3)
    add_body(
        doc,
        "Build에서 1.0.0 (2)를 고른다. "
        "빌드 1은 위젯 없음, 빌드 2는 위젯 있음. "
        "아카이브 스킴은 SubscriptionTracker. "
        "SubscriptionWidgetExtension으로 잡혀 있으면 위젯만 올리려 한다.",
    )

    add_heading_styled(doc, "5.6 App Review Information", 3)
    add_body(
        doc,
        "Sign-in required는 끈다. 켜 두면 심사자가 아이디를 찾고 없으면 반려된다.",
    )
    make_table(
        doc,
        ["칸", "이번에 넣은 값"],
        [
            ["First Name", "DONGMIN"],
            ["Last Name", "HWANG"],
            ["Phone", "받을 수 있는 한국 번호. +82로 시작"],
            ["Email", "hdmin8712@gmail.com"],
        ],
        [4.0, 12.6],
    )
    add_body(doc, "Notes에 넣은 내용", before=8)
    add_body(
        doc,
        "로그인 없는 앱입니다. 계정 없이 바로 사용할 수 있습니다.\n\n"
        "알림을 확인하려면:\n"
        "1. 설정 탭에서 알림 받기를 켭니다.\n"
        "2. 구독을 하나 추가하고, 다음 결제일을 내일로 둡니다.\n"
        "3. 알림 시각이 지나기 전이면 예약만 되고, "
        "결제일 하루 전 지정한 시각에 알림이 옵니다.\n\n"
        "홈 화면 위젯(작은 것, 중간 것)을 추가하면 이번 달 구독 금액을 볼 수 있습니다.",
        before=4,
    )
    add_body(
        doc,
        "Version Release는 Manually release this version.",
        before=8,
    )

    add_heading_styled(doc, "6. App Privacy", 2)
    add_body(
        doc,
        "처리방침 URL만으로는 부족하다. Get Started로 설문을 하고 Publish를 눌러야 한다. "
        "이 앱은 서버로 아무것도 안 보낸다. "
        "No, we do not collect data / Data Not Collected. "
        "Publish를 안 누르면 스토어에도 안 붙고 제출도 막힐 수 있다.",
    )

    add_heading_styled(doc, "7. Pricing and Availability", 2)
    add_body(
        doc,
        "Paid Applications 계약이 없어도 무료는 설정된다. 유료로 두면 막힌다.",
    )
    add_number(doc, "Add Pricing")
    add_number(
        doc,
        "Base Country는 Korea, Republic of (KRW). "
        "Price 목록 맨 위 ₩0.00. Free라는 글자는 없다.",
    )
    add_number(doc, "다음 화면에서 다른 나라도 0.00인지 보고 Next")
    add_number(doc, "Set Up Availability → All Countries or Regions. Pre-Order 아님")
    add_body(
        doc,
        "저장 후 각 나라 상태가 Available on App Release인 것은 정상이다. "
        "출시 전이므로 기다릴 필요 없다.",
        before=8,
    )

    add_heading_styled(doc, "8. 제출 버튼", 2)
    add_number(doc, "버전 페이지에서 Save")
    add_number(doc, "Add for Review")
    add_number(
        doc,
        "Draft Submission에 iOS App 1.0.0 (2)가 있으면 Submit for Review. "
        "빨간 빼기는 누르지 않음",
    )
    add_number(
        doc,
        "1 Item Submitted. 심사는 최대 48시간 안내. 결과는 메일",
    )
    add_body(
        doc,
        "수출 규정·광고 식별자 창이 더 뜰 수 있다. "
        "이 앱은 암호화 예외 수준이고 광고 SDK가 없다. "
        "이번 제출에서는 그 창이 안 뜨고 바로 접수됐다.",
        before=8,
    )

    add_heading_styled(doc, "9. 심사 중 / 승인 뒤", 2)
    add_bullet(doc, "In Review 동안 빌드와 스토어 글은 바꾸지 않는다. 바꾸면 심사가 리셋될 수 있다.")
    add_bullet(doc, "거절이면 Resolution Center 사유를 보고 고친 뒤 다시 제출한다.")
    add_bullet(
        doc,
        "Approved 다음에 Release This Version. "
        "스토어 반영에 몇 시간이 걸릴 수 있다.",
    )

    add_heading_styled(doc, "10. 처리방침 페이지를 다시 올릴 때", 2)
    add_bullet(doc, "파일: docs/privacy.html")
    add_bullet(doc, "GitHub 저장소 SubscriptionTracker는 Public")
    add_bullet(doc, "Pages: main, 폴더 / (root). /docs 가 아님")
    add_bullet(
        doc,
        "확인 주소: https://dongminhwan.github.io/SubscriptionTracker/docs/privacy.html",
    )

    add_heading_styled(doc, "11. 빌드를 다시 올릴 때", 2)
    add_bullet(doc, "Xcode 스킴: SubscriptionTracker, 기기 Any iOS Device")
    add_bullet(doc, "Product → Archive → Distribute App → App Store Connect → Upload")
    add_bullet(doc, "처리가 끝나면 TestFlight에 빌드가 보인다. 버전 페이지 Build에서 고른다")
    add_bullet(doc, "내부 테스트 그룹만 쓴다. External Testing은 베타 심사")

    os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
    doc.save(OUT_PATH)


if __name__ == "__main__":
    build()
    print(OUT_PATH)
