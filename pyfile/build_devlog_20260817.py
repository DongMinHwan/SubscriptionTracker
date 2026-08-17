#!/usr/bin/env python3
"""2026-08-17 개발일지 .docx 생성."""

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

from build_docx import (
    add_body,
    add_bullet,
    add_heading_styled,
    add_number,
    configure_document,
    make_table,
    set_paragraph_format,
    style_run,
)

BASE = "/Users/hwangdongmin/Projects/NewProject_1"
OUT_PATH = f"{BASE}/개발일지/2026-08-17_개발일지.docx"
IMAGE_DIR = f"{BASE}/개발일지/이미지"

KV = [3.2, 13.4]


def add_picture(doc, filename, width_cm, caption):
    p = doc.add_paragraph()
    set_paragraph_format(p, before=8, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
    p.add_run().add_picture(os.path.join(IMAGE_DIR, filename), width=Cm(width_cm))

    cap = doc.add_paragraph()
    set_paragraph_format(cap, before=0, after=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    style_run(cap.add_run(caption), 10, False)


def build():
    doc = Document()
    configure_document(doc)

    add_heading_styled(doc, "개발일지 — 2026년 8월 17일 (월)", 1)

    make_table(
        doc,
        ["항목", "내용"],
        [
            ["프로젝트", "구독 트래커 (학습용 첫 iOS 앱)"],
            ["작업 시간", "09:49 ~ 15:44"],
            [
                "오늘 목표",
                "기획서를 코드로 옮긴다. 화면 3개를 실제로 돌려 본다.",
            ],
            [
                "결과",
                "v1.0.0 화면 3개 구현 완료. 예정에 없던 v1.0.1(런치스크린·앱 아이콘)까지 진행.",
            ],
            [
                "한 줄 소감",
                "기획서를 먼저 써 두니 구현 중에 “이건 넣을까 말까”로 고민한 시간이 거의 없었다.",
            ],
        ],
        KV,
    )

    add_heading_styled(doc, "1. 오늘의 흐름", 2)
    make_table(
        doc,
        ["시각", "한 일", "산출물"],
        [
            ["09:49", "학습용 구독 트래커 기획서 초안 작성", "기획서/초안/*.docx"],
            [
                "~10:30",
                "v1 범위 확정. 백엔드는 이번에 안 함(나중에 GCP 검토)",
                "결정 사항",
            ],
            ["10:47", "Xcode 프로젝트 생성", "SubscriptionTracker.xcodeproj"],
            [
                "10:54 ~ 10:56",
                "SwiftData 모델, 유틸, 디자인 토큰, 화면 3개 구현",
                "Swift 파일 9개",
            ],
            [
                "~11:00",
                "시뮬레이터에서 실제 데이터 입력해 확인 (Cursor Pro+ / 97,000원 / 9월 17일)",
                "동작 확인",
            ],
            ["15:13", "목록 화면 스크롤 동작 수정", "SubscriptionListView.swift"],
            ["15:19", "기획서를 버전 폴더로 정리", "기획서/v1.0.0/"],
            [
                "15:33",
                "런치스크린·앱 아이콘 기획서 작성",
                "기획서/v1.0.1/*.docx",
            ],
            [
                "15:39 ~ 15:44",
                "아이콘·런치스크린 에셋 제작 및 적용, 시뮬레이터 검증",
                "Assets.xcassets, Info.plist",
            ],
        ],
        [2.6, 9.4, 4.6],
    )

    add_heading_styled(doc, "2. 기획 — 무엇을 안 만들지 정했다", 2)
    add_body(
        doc,
        "기획을 해 본 적이 없어서 제일 막막했던 부분이다. 오늘 해 보니 기획은 "
        "“무엇을 만들지” 적는 일이라기보다 “무엇을 안 만들지” 적는 일에 가까웠다.",
    )
    add_body(doc, "확정한 것:", before=6, after=4)
    add_bullet(doc, "스택은 iOS 네이티브(SwiftUI) 유지. 크로스플랫폼은 Android가 필요해지면 그때 다시 본다.")
    add_bullet(doc, "첫 앱은 학습용. 기획–디자인–개발–출시–운영을 한 바퀴 도는 것이 목적이다.")
    add_bullet(doc, "v1.0.0은 화면 3개(목록·추가·상세), 필드 4개(이름·금액·주기·다음 결제일).")
    add_bullet(doc, "백엔드 없음. SwiftData로 기기에만 저장한다.")
    add_bullet(doc, "알림, 위젯, 카테고리, 통계, 로그인은 전부 v1.0.0 비범위.")
    add_body(
        doc,
        "특히 알림과 위젯을 뺀 게 컸다. 처음엔 넣고 싶었는데, 빼고 나니 하루 만에 화면 3개가 돌아갔다. "
        "넣었으면 오늘 안에 못 끝냈을 것이다.",
        before=6,
    )

    add_heading_styled(doc, "3. 개발 — 화면 3개", 2)
    make_table(
        doc,
        ["파일", "역할"],
        [
            ["Subscription.swift", "SwiftData 모델. 결제 주기, 다음 결제일 넘김 로직"],
            ["MonthlyTotal.swift", "이번 달 합계 계산 규칙"],
            ["Formatters.swift", "금액(₩17,000), 날짜(3월 21일) 표기"],
            ["DesignTokens.swift", "기획서 11장 토큰(색·글꼴·치수)을 코드로 옮긴 것"],
            ["SubscriptionListView.swift", "목록 화면. 이번 달 총액, 빈 상태"],
            ["AddSubscriptionView.swift", "추가 시트"],
            ["SubscriptionDetailView.swift", "상세·수정·삭제"],
            ["PreviewSupport.swift", "프리뷰용 인메모리 컨테이너와 샘플 데이터"],
        ],
        [6.0, 10.6],
    )
    add_body(
        doc,
        "템플릿으로 생긴 Item.swift와 ContentView.swift는 지웠다.",
        before=8,
    )
    add_body(
        doc,
        "색과 치수를 DesignTokens 한 곳에 모아 둔 게 뒤에서 효과를 봤다. "
        "화면마다 색을 직접 적어 뒀다면 v1.0.1에서 런치스크린 배경을 맞출 때 "
        "어디를 봐야 할지부터 헤맸을 것이다.",
        before=6,
    )
    add_picture(doc, "2026-08-17_목록화면.png", 5.2, "목록 화면 (빈 상태)")

    add_heading_styled(doc, "4. 고친 것 — 스크롤이 어색했다", 2)
    add_body(
        doc,
        "직접 써 보니 “구독” 타이틀과 총액이 한 덩어리로 같이 움직여서 어색했다. "
        "원한 동작은 타이틀은 고정, 총액은 리스트와 같이 스크롤되는 것이었다.",
    )
    add_body(doc, "고친 방법:", before=6, after=4)
    add_number(doc, "타이틀을 navigationBarTitleDisplayMode(.inline)로 내비게이션 바에 고정했다.")
    add_number(doc, "총액 영역을 List 바깥이 아니라 List의 첫 번째 행으로 옮겼다.")
    add_number(doc, "그 행만 listRowInsets와 배경을 조정해 리스트 행처럼 보이지 않게 했다.")
    add_body(
        doc,
        "기획서에 “총액은 목록과 함께 스크롤된다”라고 한 줄만 적혀 있었는데, "
        "그 한 줄을 List 안에 넣느냐 밖에 두느냐로 결과가 완전히 달라졌다. "
        "기획서의 문장이 곧 구현 구조라는 걸 처음 체감했다.",
        before=6,
    )

    add_heading_styled(doc, "5. v1.0.1 — 런치스크린과 앱 아이콘", 2)
    add_body(
        doc,
        "앱을 켤 때 흰 화면이 잠깐 뜨고, 홈 화면에는 회색 사각형이 보였다. "
        "기능이 아니라 첫인상 문제라 별도 버전으로 떼어 기획서를 따로 썼다.",
    )
    make_table(
        doc,
        ["항목", "내용"],
        [
            ["런치스크린", "배경 #F2F2F7 + 가운데 로고 120pt. 정적 1장."],
            ["앱 아이콘", "1024×1024, 배경 #1F4E79, 흰 마크, 알파 없음."],
            ["로고 마크", "순환 화살표 + 원화 기호. 매달 반복해서 나가는 돈이라는 뜻."],
            ["추가 토큰", "launch.background, launch.logoSize, icon.background, icon.mark"],
        ],
        KV,
    )
    add_body(
        doc,
        "런치스크린은 스플래시가 아니다. 앱 코드가 실행되기 전에 시스템이 그리는 정적 화면이라 "
        "애니메이션도, 이번 달 총액 같은 데이터도 넣을 수 없다. "
        "그래서 첫 화면(목록)과 같은 배경색을 써서 자연스럽게 이어지게만 했다.",
        before=8,
    )
    add_picture(doc, "2026-08-17_런치스크린.png", 5.2, "런치스크린")
    add_picture(doc, "2026-08-17_앱아이콘.png", 3.6, "홈 화면의 앱 아이콘")

    add_heading_styled(doc, "6. 오늘 막힌 것", 2)

    add_heading_styled(doc, "6.1 없는 빌드 설정을 한참 붙잡고 있었다", 3)
    add_body(
        doc,
        "런치스크린 배경색과 로고를 INFOPLIST_KEY_UILaunchScreen_UIColorName / _UIImageName "
        "빌드 설정으로 넣으려 했다. 빌드는 성공하는데 생성된 Info.plist에는 아무것도 안 들어갔다. "
        "오류도 경고도 없어서 원인을 찾는 데 시간이 걸렸다.",
    )
    add_body(
        doc,
        "Xcode의 빌드 설정 정의 파일(CoreBuildSystem.xcspec)을 직접 열어 보고 알았다. "
        "런치스크린 관련 설정은 INFOPLIST_KEY_UILaunchScreen_Generation 하나뿐이고, "
        "설명이 “UILaunchScreen 키를 빈 딕셔너리로 설정한다”였다. "
        "지금까지 흰 화면이 떴던 원인이 바로 이거였다.",
        before=6,
    )
    add_body(
        doc,
        "해결: Info.plist 파일을 직접 만들고 INFOPLIST_FILE로 지정했다. "
        "GENERATE_INFOPLIST_FILE = YES를 그대로 두면 Xcode가 이 파일 위에 나머지 키를 얹어 준다. "
        "둘 중 하나만 써야 하는 줄 알았는데 병합된다는 걸 오늘 알았다.",
        before=6,
    )

    add_heading_styled(doc, "6.2 아이콘 배경색이 기획서와 달랐다", 3)
    add_body(
        doc,
        "생성한 로고 그림에 옅은 비네트(가장자리가 어두워지는 효과)가 있어서, "
        "그대로 쓰면 배경이 #1F4E79가 아니었다. 기획서에 “단색 평면, 그라데이션 없음”이라고 적어 둔 것과 어긋났다.",
    )
    add_body(
        doc,
        "밝기 분포를 보니 배경은 58~63, 마크는 248~255로 뚜렷하게 갈렸다. "
        "그 사이 값만 가장자리로 처리하고 나머지는 눌러서, 마크 모양만 뽑아 각각 다시 칠했다. "
        "덕분에 아이콘과 런치스크린의 마크가 완전히 같고 색도 토큰과 정확히 일치한다.",
        before=6,
    )

    add_heading_styled(doc, "6.3 되돌릴 방법이 없었다", 3)
    add_body(
        doc,
        "런치스크린 때문에 project.pbxproj를 건드려야 했는데, 이 프로젝트가 아직 git 저장소가 아니라 "
        "되돌릴 방법이 없었다. 급한 대로 .bak 파일로 복사해 두고 작업했다. "
        "오늘 가장 아찔했던 부분이고, 내일 제일 먼저 할 일이 됐다.",
    )

    add_heading_styled(doc, "7. 오늘 배운 것", 2)
    add_number(
        doc,
        "INFOPLIST_KEY_* 는 아무 Info.plist 키나 되는 게 아니다. Xcode가 정의해 둔 목록만 동작하고, "
        "없는 이름을 써도 조용히 무시된다.",
    )
    add_number(
        doc,
        "GENERATE_INFOPLIST_FILE 과 INFOPLIST_FILE 은 배타적이지 않다. 같이 쓰면 병합된다.",
    )
    add_number(
        doc,
        "런치스크린은 스플래시 화면이 아니다. 애니메이션도 데이터도 못 넣는 정적 화면이다.",
    )
    add_number(
        doc,
        "시뮬레이터에서 런치스크린처럼 순식간에 지나가는 화면은 "
        "simctl launch --wait-for-debugger 로 앱을 멈춰 세우면 천천히 확인할 수 있다.",
    )
    add_number(
        doc,
        "기획서에 범위를 적는 것보다 비범위를 적는 게 실제로는 더 도움이 됐다.",
    )

    add_heading_styled(doc, "8. 남은 것 / 다음에 할 일", 2)
    make_table(
        doc,
        ["순서", "할 일", "이유"],
        [
            [
                "1",
                "git 저장소 초기화하고 오늘 작업을 첫 커밋으로 남기기",
                "지금은 실수하면 되돌릴 방법이 없다",
            ],
            [
                "2",
                "project.pbxproj.bak-before-v1.0.1 확인 후 정리",
                "git이 생기면 필요 없어진다",
            ],
            [
                "3",
                "MARKETING_VERSION 을 1.0.1 로 올릴지 결정",
                "기획서 범위 표에 없어서 아직 1.0 그대로 둠",
            ],
            [
                "4",
                "실제 기기에서 아이콘·런치스크린 확인",
                "오늘은 시뮬레이터에서만 봤다",
            ],
            [
                "5",
                "출시 준비 알아보기 (App Store Connect, 스크린샷, 개인정보 처리방침)",
                "학습 목적이 출시까지 가 보는 것이므로",
            ],
        ],
        [1.6, 8.0, 7.0],
    )
    add_body(
        doc,
        "알림과 위젯은 계속 v1.0.0 비범위로 둔다. 출시를 한 번 겪어 본 다음에 다음 버전으로 뺀다.",
        before=8,
    )

    add_heading_styled(doc, "9. 오늘 만진 파일", 2)
    make_table(
        doc,
        ["구분", "경로"],
        [
            [
                "새로 만듦",
                "SubscriptionTracker/SubscriptionTracker/*.swift (9개)\n"
                "SubscriptionTracker/Info.plist\n"
                "Assets.xcassets/AppIcon.appiconset/AppIcon-1024.png\n"
                "Assets.xcassets/LaunchLogo.imageset/ (@1x·@2x·@3x)\n"
                "Assets.xcassets/LaunchBackground.colorset/\n"
                "기획서/v1.0.0/, 기획서/v1.0.1/\n"
                "pyfile/build_launch_assets.py",
            ],
            [
                "고침",
                "SubscriptionListView.swift (스크롤)\n"
                "DesignTokens.swift (토큰 4개 추가)\n"
                "SubscriptionTracker.xcodeproj/project.pbxproj (INFOPLIST_FILE)",
            ],
            ["지움", "Item.swift, ContentView.swift (템플릿 파일)"],
        ],
        [2.6, 14.0],
    )

    os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
    doc.save(OUT_PATH)


if __name__ == "__main__":
    build()
    print(OUT_PATH)
