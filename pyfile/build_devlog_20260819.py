#!/usr/bin/env python3
"""2026-08-19 개발일지 .docx 생성."""

import os

from build_docx import (
    add_body,
    add_bullet,
    add_heading_styled,
    add_number,
    configure_document,
    make_table,
)
from docx import Document

BASE = "/Users/hwangdongmin/Projects/NewProject_1"
OUT_PATH = f"{BASE}/개발일지/2026-08-19_개발일지.docx"

KV = [3.2, 13.4]


def build():
    doc = Document()
    configure_document(doc)

    add_heading_styled(doc, "개발일지 — 2026년 8월 19일 (수)", 1)

    make_table(
        doc,
        ["항목", "내용"],
        [
            ["날짜", "2026년 8월 19일 저녁"],
            ["프로젝트", "구독 트래커 (학습용 첫 iOS 앱)"],
            [
                "목표",
                "심사를 올리지 않고, v1.1.0 방향을 다시 잡은 뒤 "
                "실기기로 한 바퀴 돌려 본다.",
            ],
            [
                "결과",
                "위젯은 보류로 남겼다. 탭·아이콘·알림으로 v1.1.0을 다시 썼다. "
                "총액이 세는 것과 그 위 말을 고쳤다. "
                "App Store Connect에 “내 구독 트래커”로 등록하고 "
                "1.0.0 (1)을 TestFlight로 실기기에 설치했다.",
            ],
            [
                "한 줄 소감",
                "심사 준비가 아니라 받을 길을 먼저 열었다. "
                "이름은 홈 화면과 스토어가 달라도 되고, 숫자는 말이 맞아야 한다.",
            ],
        ],
        KV,
    )

    add_heading_styled(doc, "1. 오늘의 흐름", 2)
    make_table(
        doc,
        ["시각", "한 일", "결과"],
        [
            [
                "20:30 ~ 21:00",
                "개발자 계약 확인, 심사 범위 논의",
                "위젯 없이 내고 앱을 꾸미자는 쪽에서 출발",
            ],
            [
                "21:00 ~ 21:20",
                "v1.1.0 범위를 다시 잡음",
                "탭 2개 + 아이콘 + 알림. 통계·로고·위젯은 다음",
            ],
            [
                "21:20 ~ 21:40",
                "아이콘 시안·구현, 색이 겹치는 문제 수정",
                "라운드 사각, 가장 적게 쓰인 색을 줌",
            ],
            [
                "21:40 ~ 22:10",
                "날짜 영어, 총액 정의, 라벨, 탭·알림",
                "ko, 8월 구독 금액 ₩14,000, 구독/설정 탭",
            ],
            [
                "21:44 ~ 22:29",
                "앱 등록 → 아카이브 → TestFlight",
                "실기기에서 설치·실행 확인",
            ],
        ],
        [3.4, 7.6, 5.6],
    )
    add_body(
        doc,
        "시각은 대화와 스크린샷 기준으로 대략 적은 것이다. "
        "저녁에 기획을 다시 잡고, 구현하다가, 밤에는 스토어 쪽으로 넘어갔다.",
        before=8,
    )

    add_heading_styled(doc, "2. 심사 대신 TestFlight를 먼저 연 이유", 2)
    add_body(
        doc,
        "개발자 프로그램 계약은 오늘 수락됐다. Paid Applications는 무료 앱이라 해당 없다. "
        "심사를 바로 올릴 여건은 갖춰졌지만, 위젯을 막아 둔 채로 화면 3개만 내면 "
        "가이드라인 4.2에 대한 답이 약하다.",
    )
    add_body(
        doc,
        "꾸미기만으로는 답이 안 된다. 색과 애니메이션은 “유용하거나, 고유하거나, 앱다운 것”에 "
        "들어가지 않는다. 그래서 알림을 넣고, 알림이 들어갈 자리로 설정 탭을 두기로 했다.",
        before=6,
    )
    add_body(
        doc,
        "오늘은 심사를 올리지 않기로 했다. 앱 자리를 만들고 아카이브해서 TestFlight로 "
        "실기기에 받아 보는 것까지만 했다. 받을 길이 열리면 고친 것을 바로 확인할 수 있다.",
        before=6,
    )

    add_heading_styled(doc, "3. v1.1.0을 다시 씀", 2)
    add_body(
        doc,
        "원래 v1.1.0은 홈 화면 위젯이었다. 어젯밤 “이번 달” 정의가 흔들려 막아 뒀고, "
        "오늘은 그 자리를 탭·아이콘·알림이 대신한다.",
    )
    add_body(
        doc,
        "위젯 기획서는 지우지 않았다. 기획서/보류/위젯으로 옮겼다. "
        "코드도 임베드 한 줄만 빼 둔 상태 그대로다. 폐기가 아니라 보류다.",
        before=6,
    )

    add_heading_styled(doc, "3.1 하는 것 / 하지 않는 것", 3)
    make_table(
        doc,
        ["하는 것", "하지 않는 것"],
        [
            [
                "탭 2개 (구독, 설정)\n"
                "구독 아이콘 (머리글자 + 색)\n"
                "결제 전날 알림\n"
                "설정 화면\n"
                "총액 판정과 라벨 수정",
                "위젯 (보류)\n"
                "통계·차트 (다음 버전)\n"
                "결제 이력 (다음 버전)\n"
                "브랜드 로고\n"
                "로그인·서버",
            ]
        ],
        [8.3, 8.3],
    )

    add_heading_styled(doc, "3.2 브랜드 로고를 안 쓰기로 한 이유", 3)
    add_bullet(
        doc,
        "권리. 회사마다 브랜드 자산 조건이 다르고, 앱에 넣어 배포하는 것은 대개 별도 승인이다. "
        "가이드라인 5.2도 권리 없는 상표를 반려 사유로 둔다.",
    )
    add_bullet(
        doc,
        "빈칸. 로고를 자동으로 받아와도 네이버멤버십, 쿠팡와우, 밀리의서재는 안 나온다. "
        "넷플릭스만 뜨면 미완성으로 보인다.",
    )
    add_body(
        doc,
        "머리글자와 색은 어떤 이름에도 통하고, 권리 문제가 없고, 사용자가 입력할 것도 없다.",
        before=6,
    )

    add_heading_styled(doc, "4. 아이콘", 2)
    add_body(
        doc,
        "네 가지를 그렸다. A 없음, B 채운 원, C 옅은 원, D 라운드 사각. D로 정했다. "
        "홈 화면 앱 아이콘과 같은 모양이라 “이 서비스”로 읽히고, 나중에 실제 로고를 넣게 되면 "
        "같은 자리에 같은 모양으로 바꿔 끼울 수 있다.",
    )

    add_heading_styled(doc, "4.1 색이 겹쳤다", 3)
    add_body(
        doc,
        "처음에는 이름을 해시해서 색을 정했다. 넣어 보니 넷플릭스와 애플 개발 프로그램이 "
        "둘 다 주황이었다. 여덟 칸에 무작위로 떨어뜨리면 둘만 있어도 여덟 번에 한 번은 겹친다. "
        "색을 넣은 이유가 구분인데 구분이 안 됐다.",
    )
    add_body(
        doc,
        "지금은 목록에서 가장 적게 쓰인 색을 준다. 같으면 앞 번호부터다. "
        "여덟 개까지는 안 겹친다. 예전에 만든 구독은 앱을 열 때 한 번 채운다.",
        before=6,
    )
    add_body(
        doc,
        "추가 화면에서 이름이 비어 있을 때 물음표를 띄우던 것도 뺐다. "
        "아직 안 친 것인데 잘못된 것처럼 보였다. 색만 둔다.",
        before=6,
    )

    add_heading_styled(doc, "5. 앱이 영어였다", 2)
    add_body(
        doc,
        "추가 화면 날짜가 “Aug 19, 2026”으로 나왔다. 화면 글자는 한국어로 박아 넣어 티가 안 났는데, "
        "날짜 선택기처럼 시스템이 그리는 부분은 앱의 언어를 따른다.",
    )
    add_body(
        doc,
        "원인: project.pbxproj의 developmentRegion이 en. 지원 언어가 영어뿐이라 "
        "한국어 아이폰에서도 영어로 나오고, 이대로면 App Store에도 영어 앱으로 등록된다. "
        "ko로 바꿨다. 달력은 “2026년 8월”로 나왔다.",
        before=6,
    )

    add_heading_styled(doc, "6. 총액 — 숨어 있던 정의 문제를 오늘 정함", 2)
    add_body(
        doc,
        "어젯밤 위젯을 막은 이유가 오늘 목록 위에서도 그대로였다. "
        "구독이 둘인데 총액이 ₩143,000. 넷플릭스 ₩14,000 + 애플 개발 프로그램 ₩129,000. "
        "그런데 애플 개발 프로그램은 D-364, 내년 8월에 나갈 돈이다.",
    )
    add_body(
        doc,
        "매년 구독을 판정할 때 연도를 안 보고 월만 비교한 탓이다. "
        "2027년 8월과 2026년 8월이 둘 다 “8월”이라 같은 달로 처리됐다.",
        before=6,
    )

    add_heading_styled(doc, "6.1 정한 규칙", 3)
    add_body(
        doc,
        "다음 결제일이 이번 달 안에 있으면 더한다. 연도와 월을 함께 본다. "
        "매달인지 매년인지는 보지 않는다. 주기를 나눠 따질 필요가 없어졌다.",
    )
    make_table(
        doc,
        ["오늘", "구독", "다음 결제일", "포함"],
        [
            ["8월 19일", "넷플릭스", "8월 28일 (D-9)", "○"],
            ["8월 19일", "애플 개발 프로그램", "2027년 8월 18일 (D-364)", "×"],
            ["8월 30일", "넷플릭스", "9월 8일 (D-9)", "×"],
        ],
        [3.0, 5.2, 5.6, 2.8],
    )
    add_body(
        doc,
        "고친 뒤 총액은 ₩14,000이 됐다. 이게 맞다.",
        before=8,
    )

    add_heading_styled(doc, "6.2 라벨", 3)
    add_body(
        doc,
        "“이번 달”은 숨기기로 했다가, 금액만 남으면 무엇을 더한 숫자인지 모른다. "
        "후보를 놓고 “8월 구독 금액”으로 정했다. 달이 바뀌면 숫자도 따라 바뀐다.",
    )
    add_body(
        doc,
        "받아들여야 할 결과가 있다. 결제일이 지나면 다음 주기로 밀리므로 "
        "총액은 달이 갈수록 줄고, 그 달 결제가 다 끝나면 월말 며칠 동안 ₩0이 된다. "
        "이미 나간 돈까지 세려면 결제 이력이 필요하고, 그건 다음 버전이다.",
        before=6,
    )

    add_heading_styled(doc, "7. 탭·설정·알림", 2)
    add_body(
        doc,
        "탭만 나누면 기능이 늘지 않는다. 알림이 들어가면서 설정할 것이 생긴다. "
        "알림을 받을지, 전날 몇 시에 받을지. 그래서 탭과 알림은 한 묶음이다.",
    )
    add_bullet(doc, "탭: 구독 / 설정. 기본은 구독.")
    add_bullet(doc, "알림: 결제일 하루 전, 기본 오전 9시. 같은 날 여러 건이면 하루에 하나.")
    add_bullet(doc, "권한: 앱 첫 실행이 아니라 설정에서 스위치를 켤 때 묻는다.")
    add_bullet(doc, "예약: 앞으로 60일 안. iOS 한도가 64개라 넘으면 조용히 잘린다.")
    add_bullet(doc, "구독이 바뀌면 예약을 전부 지우고 다시 건다.")

    add_heading_styled(doc, "8. 앱 등록과 이름", 2)
    add_body(
        doc,
        "아카이브보다 App Store Connect에 앱 자리를 먼저 만들어야 한다. "
        "플랫폼 iOS, 주 언어 Korean, 번들 ID는 앱 본체 "
        "(com.dm.SubscriptionTracker). 위젯 번들 ID는 고르면 안 된다.",
    )
    add_body(
        doc,
        "SKU는 subscription-tracker. User Access는 Full Access. "
        "혼자 쓰는 계정이라 제한할 사람이 없다.",
        before=6,
    )

    add_heading_styled(doc, "8.1 스토어 이름이 이미 쓰이고 있었다", 3)
    add_body(
        doc,
        "Name에 “구독 트래커”를 넣으니 "
        "“The app name you entered is already being used.” 가 나왔다. "
        "설명하는 이름이라 상표 주장은 하지 않았다.",
    )
    add_body(
        doc,
        "스토어 이름과 홈 화면 이름은 따로다. 홈 화면은 그대로 “구독 트래커”로 둬도 된다. "
        "MySubTracker도 후보였으나 앱은 한국어인데 스토어만 영어가 된다. "
        "“내 구독 트래커”로 등록했다. Create가 통과했다.",
        before=6,
    )

    add_heading_styled(doc, "9. TestFlight", 2)
    add_body(
        doc,
        "Xcode Organizer에서 아카이브 후 업로드. "
        "SubscriptionTracker 1.0.0 (1) uploaded. 처리가 끝나면 TestFlight 탭에 빌드가 보인다.",
    )
    add_body(
        doc,
        "All Testers 화면에서는 사람을 넣을 수 없다. 이미 초대한 사람을 모아 보는 곳이다. "
        "Internal Testing 그룹을 만들고(이름은 Test) 본인(sizzle@naver.com)을 넣었다. "
        "External Testing은 베타 심사가 붙으므로 건드리지 않았다.",
        before=6,
    )
    add_body(
        doc,
        "처음에는 No Builds Available. 그룹 Builds 탭에는 이미 1.0.0 (1) Ready to Test가 붙어 있었다. "
        "맥에서 더 넣을 것은 없고, 아이폰 TestFlight를 같은 Apple ID로 열면 된다. "
        "앱이 보였고, 설치해서 실행까지 확인했다.",
        before=6,
    )
    add_body(
        doc,
        "iOS Builds의 노란 Ready to Submit은 스토어 심사 제출용 표시다. "
        "내부 테스트와는 별개라 지금은 누르지 않는다.",
        before=6,
    )

    add_heading_styled(doc, "10. 오늘 배운 것", 2)
    add_number(
        doc,
        "받을 길을 먼저 열면 시뮬레이터와 실기기의 간극을 바로 본다. "
        "심사 제출과 아카이브는 다른 일이다.",
    )
    add_number(
        doc,
        "스토어 이름은 홈 화면 이름과 달라도 된다. 다만 언어는 맞추는 편이 낫다. "
        "이미 쓰이는 설명 이름에 상표 주장은 하지 말고 한 단어를 붙인다.",
    )
    add_number(
        doc,
        "번들 ID 드롭다운에 위젯이 같이 나온다. 앱 자리를 만들 때는 본체만 고른다.",
    )
    add_number(
        doc,
        "화면 글자를 한국어로 박아 넣어도 앱 언어가 en이면 시스템이 그리는 부분은 영어다. "
        "developmentRegion을 봐야 한다.",
    )
    add_number(
        doc,
        "색을 구분용으로 넣을 거면 해시보다 “안 겹치게”가 먼저다. "
        "여덟 칸에 둘을 던지면 겹친다.",
    )
    add_number(
        doc,
        "월만 비교하면 내년 같은 달이 이번 달에 들어온다. 연도를 같이 봐야 한다.",
    )
    add_number(
        doc,
        "탭은 배치지 기능이 아니다. 두 번째 탭이 성립하려면 그 안에 담을 것이 먼저 있어야 한다.",
    )

    add_heading_styled(doc, "11. 남은 것 / 다음에 할 일", 2)
    make_table(
        doc,
        ["순서", "할 일", "메모"],
        [
            [
                "1",
                "오늘 작업을 GitHub에 저장한다",
                "원격은 이미 있다. 커밋만 안 올라간 상태",
            ],
            [
                "2",
                "알림이 실기기에서 오는지 확인한다",
                "설정에서 스위치를 켜고 권한을 허용한 뒤, 결제일을 내일로 둔 구독으로",
            ],
            [
                "3",
                "MARKETING_VERSION을 1.1.0으로 올린다",
                "지금 올린 빌드는 1.0.0 (1). 다음 빌드부터",
            ],
            [
                "4",
                "개인정보 처리방침 주소",
                "App Store Connect 필수. 설정 화면 링크도 여기",
            ],
            [
                "5",
                "스크린샷·설명·키워드",
                "심사 제출 직전에. 지금은 TestFlight만",
            ],
            [
                "6",
                "위젯 되살리기",
                "총액 정의는 오늘 정했다. 임베드 한 줄. 다음 버전에서",
            ],
            [
                "7",
                "결제 이력 → 통계",
                "월말에 ₩0이 되는 한계를 풀려면 이력이 필요하다",
            ],
        ],
        [1.6, 6.6, 8.4],
    )

    add_heading_styled(doc, "12. 오늘 만진 파일", 2)
    make_table(
        doc,
        ["파일", "무엇을"],
        [
            ["기획서/보류/위젯/", "옛 v1.1.0 위젯 기획서를 옮김"],
            ["기획서/v1.1.0/", "탭·아이콘·알림 기획서와 시안"],
            ["Shared/Subscription.swift", "colorIndex 추가"],
            ["Shared/SubscriptionIcon.swift", "머리글자, 색, 아이콘 뷰"],
            ["Shared/DesignTokens.swift", "아이콘 색·크기 토큰"],
            ["Shared/MonthlyTotal.swift", "연도까지 보는 포함 판정"],
            ["Shared/Formatters.swift", "8월 구독 금액 라벨"],
            ["SubscriptionListView.swift", "아이콘, 색 채우기, 라벨"],
            ["Add/Detail + IconColorPicker", "색 고르기"],
            ["RootView.swift", "탭 2개"],
            ["SettingsView.swift", "알림 스위치·시각, 버전"],
            ["PaymentNotifications.swift", "전날 알림 예약"],
            ["WidgetRefresh.swift", "저장 후 위젯·알림 같이 맞춤"],
            ["project.pbxproj", "developmentRegion = ko"],
        ],
        [7.2, 9.4],
    )

    add_heading_styled(doc, "13. 확인에 쓴 방법", 2)
    add_bullet(doc, "시뮬레이터 빌드·설치 후 스크린샷으로 아이콘·총액·라벨 확인")
    add_bullet(doc, "Xcode Product → Archive → Distribute App → Upload")
    add_bullet(doc, "App Store Connect → TestFlight → Internal Testing")
    add_bullet(doc, "아이폰 TestFlight에서 설치·실행")
    add_body(
        doc,
        "시뮬레이터에서 맞다고 끝난 것이 아니다. 오늘 처음으로 실기기에서 같은 앱을 열었다. "
        "다음 확인은 알림이 실제로 오는가다.",
        before=6,
    )

    os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
    doc.save(OUT_PATH)


if __name__ == "__main__":
    build()
    print(OUT_PATH)
