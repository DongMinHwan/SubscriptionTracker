#!/usr/bin/env python3
"""2026-08-18 개발일지 두 번째 .docx 생성 (위젯 작업분)."""

import os

from build_docx import (
    add_body,
    add_bullet,
    add_heading_styled,
    add_number,
    configure_document,
    make_table,
)
from docx import Document

BASE = "/Users/hwangdongmin/Projects/NewProject_1"
OUT_PATH = f"{BASE}/개발일지/2026-08-18_개발일지_2.docx"

KV = [3.2, 13.4]


def build():
    doc = Document()
    configure_document(doc)

    add_heading_styled(doc, "개발일지 — 2026년 8월 18일 (화) 두 번째", 1)

    make_table(
        doc,
        ["항목", "내용"],
        [
            ["날짜", "2026년 8월 18일 밤 ~ 19일 새벽"],
            ["프로젝트", "구독 트래커 (학습용 첫 iOS 앱)"],
            ["목표", "v1.1.0 위젯을 마무리한다."],
            [
                "결과",
                "위젯은 완성했다가 도로 막아 뒀다. 날짜 표기를 D-day로 바꿨다. "
                "“이번 달”이라는 말의 정의가 흔들린다는 것을 발견했다.",
            ],
            [
                "한 줄 소감",
                "표시 방식을 바꿨더니 숨어 있던 정의 문제가 튀어나왔다. 화면이 이상하면 대개 화면 탓이 아니다.",
            ],
        ],
        KV,
    )

    add_heading_styled(doc, "1. 오늘의 흐름", 2)
    make_table(
        doc,
        ["시각", "한 일", "결과"],
        [
            ["23:00 ~ 23:20", "작은 위젯 문구 다듬기", "“구독 3개” → “네이버맴버십 외 2개”"],
            ["23:20 ~ 23:30", "총액과 개수의 계산 기준 통일", "숨어 있던 불일치 1건 수정"],
            ["23:30 ~ 23:50", "저장 즉시 결제일 밀어내기", "넣었다 되돌렸다 다시 넣음"],
            ["23:50 ~ 00:10", "“이번 달” 총액의 정의를 기획서와 대조", "기획대로 맞음을 확인"],
            ["00:10 ~ 00:25", "날짜 표기를 D-day로 전환", "앱·위젯 동시 변경"],
            ["00:25 ~ 00:32", "위젯 임시 차단", "정의를 정할 때까지 보류"],
        ],
        [3.2, 8.4, 5.0],
    )

    add_heading_styled(doc, "2. 작은 위젯 — 숫자보다 이름", 2)
    add_body(
        doc,
        "작은 위젯 아랫줄이 “구독 3개”였다. 숫자만 있으면 무엇을 센 건지 와닿지 않아 "
        "“네이버맴버십 외 2개”로 바꿨다.",
    )
    add_body(
        doc,
        "대표 이름은 결제일이 가장 이른 것을 고른다. 그래야 중간 위젯 목록의 첫 줄과 항상 같은 이름이 나온다. "
        "두 위젯을 나란히 놓았을 때 서로 딴소리하면 안 된다.",
        before=6,
    )
    add_body(
        doc,
        "한 덩어리 문장으로 두면 이름이 길 때 “유튜브 프리미엄 플…”이 되어 개수가 통째로 날아간다. "
        "이름과 개수를 따로 두고 자를 쪽을 이름으로 못박았다. 이름 뒷부분보다 몇 개인지가 아깝다.",
        before=6,
    )

    add_heading_styled(doc, "3. 총액과 개수가 다른 기준을 쓰고 있었다", 2)
    add_body(
        doc,
        "문구를 고치다가 발견했다. 총액은 이번 달에 결제되는 것만 더하는데, 개수는 전체 구독 수를 세고 있었다.",
    )
    make_table(
        doc,
        ["상황", "총액", "개수"],
        [
            ["매달 구독 3개 + 12월 결제 연간 구독 1개 (8월에 보면)", "3개분", "4개"],
        ],
        [8.4, 4.1, 4.1],
    )
    add_body(
        doc,
        "지금 데이터로는 안 드러났다. 연간 구독을 하나 넣어야 보인다. "
        "판정을 MonthlyTotal.includes로 따로 빼서 총액과 개수가 같은 함수를 쓰게 했다.",
        before=8,
    )

    add_heading_styled(doc, "3.1 고쳤더니 다음 문제가 나왔다", 3)
    add_body(
        doc,
        "개수를 이번 달 기준으로 바꾸니, 다음 달에 결제되는 연간 구독만 가진 사람은 개수가 0이 된다. "
        "그러면 “구독을 추가하면 여기에 보입니다”가 뜬다. 이미 넣은 사람에게 안 넣었다고 말하는 셈이다.",
    )
    add_body(
        doc,
        "총액이 0원인 경우가 두 가지인데 사용자가 할 일이 서로 다르다는 것을 그제야 알았다. "
        "구독 자체가 없는지 여부를 따로 들고 다니게 하고, 후자에는 “이번 달 결제 없음”을 띄운다.",
        before=6,
    )

    add_heading_styled(doc, "4. 저장 즉시 밀어내기 — 넣었다 뺐다 다시 넣었다", 2)
    add_body(
        doc,
        "이미 쓰고 있는 구독을 등록하려고 유튜브 프리미엄의 결제일을 7월 13일로 넣었다. "
        "목록에 7월 13일이 그대로 남았다. 과거 날짜인데 밀리지 않았다.",
    )
    add_body(
        doc,
        "원인은 추가 화면이 rollForwardIfNeeded를 부르지 않아서였다. 밀어내기는 목록 화면의 onAppear와 "
        "앱 활성화 시점에만 돈다. 시트로 추가한 직후엔 둘 다 일어나지 않으니, 앱을 껐다 켤 때까지 과거 날짜로 남는다.",
        before=6,
    )
    add_body(
        doc,
        "저장할 때 밀도록 고쳤다가 한 번 되돌리고 다시 넣었다. "
        "되돌린 이유는 “다음 결제일”이라는 필드를 “당일 결제일”로 잘못 읽고 있었기 때문이다. "
        "필드 이름을 다시 확인하고 나서 원래 판단이 맞았다는 것을 알았다.",
        before=6,
    )
    add_body(
        doc,
        "이 덕분에 “구독을 추가한다”가 무슨 뜻인지 다시 생각하게 됐다. "
        "기획서의 유저 플로우 1번은 “구독을 새로 결제한 직후”만 상정했는데, "
        "앱을 처음 깔면 아무도 그러지 않는다. 이미 쓰고 있는 것 여러 개를 한꺼번에 넣는다. "
        "그때 사람이 아는 건 “매달 13일에 빠져나간다”이지 “다음 결제일이 9월 13일”이 아니다.",
        before=6,
    )

    add_heading_styled(doc, "5. 날짜를 D-day로 바꿨다", 2)
    add_body(
        doc,
        "다른 구독 트래커 앱을 보다가 정했다. 그 앱은 날짜 대신 D-4, D-334 같은 배지를 쓴다.",
    )
    add_body(
        doc,
        "“8월 28일”을 읽으려면 오늘이 며칠인지 떠올려 빼야 한다. "
        "이 앱에서 날짜를 보는 이유는 “언제 나가는지”지 “며칠인지”가 아니다. 그 뺄셈을 앱이 대신하기로 했다.",
        before=6,
    )
    make_table(
        doc,
        ["상황", "표기"],
        [
            ["결제일이 남았다", "D-9"],
            ["결제일이 오늘이다", "D-DAY"],
        ],
        [8.4, 8.2],
    )
    add_body(
        doc,
        "약점은 알고 간다. D-364는 언제인지 바로 그려지지 않는다. "
        "가까우면 D-day, 멀면 날짜로 나누는 안도 있었지만 형식을 하나로 두는 쪽을 골랐다. "
        "어떤 줄은 날짜고 어떤 줄은 D-day면 눈이 두 번 적응해야 한다.",
        before=8,
    )
    add_body(
        doc,
        "고친 곳은 DateFormat 한 군데다. 앱 목록과 중간 위젯이 이미 같은 함수를 부르고 있어서 "
        "한 곳만 바꾸니 양쪽이 같이 바뀌었다. 공용 코드를 Shared로 모아 둔 값을 여기서 처음 봤다.",
        before=6,
    )

    add_heading_styled(doc, "5.1 “매년”을 붙였다가 뺐다", 3)
    add_body(
        doc,
        "매년 구독은 금액이 훨씬 크다. 129,000원이 한 달에 나가는 줄 읽힐까 봐 “매년”을 앞에 붙였었다. "
        "날짜로 적을 때는 “매년 8월 18일”이 말이 됐다.",
    )
    add_body(
        doc,
        "D-day로 바꾸니 “매년 D-364”가 되어 어색해졌다. 빼기로 했다. "
        "빼도 잃는 것이 없다. 매달 구독은 다음 결제가 늘 한 달 안이라 D-31을 넘을 수 없으니, "
        "D-364라는 숫자 자체가 이미 매년 구독이라고 말한다.",
        before=6,
    )

    add_heading_styled(doc, "6. “이번 달”의 정의가 흔들린다", 2)
    add_body(doc, "D-day로 바꾸자마자 위젯이 이런 모양이 됐다.")
    make_table(
        doc,
        ["줄", "내용"],
        [
            ["머리글", "이번 달  ₩143,000"],
            ["1행", "넷플릭스  D-9  ₩14,000"],
            ["2행", "애플 개발 프로그램  D-364  ₩129,000"],
        ],
        [3.2, 13.4],
    )
    add_body(
        doc,
        "“이번 달”이라면서 “D-364”다. 364일 뒤에 나갈 돈이 왜 이번 달 합계에 들어 있느냐는 말이 나온다.",
        before=8,
    )
    add_body(
        doc,
        "둘 다 거짓말은 아니다. 애플 개발 프로그램은 8월 18일에 이미 나갔으므로 8월 총액에 들어가고, "
        "다음 결제는 내년 8월 18일이라 D-364다. 총액은 뒤를 보고 날짜는 앞을 본다. "
        "시제가 다른 두 숫자가 한 줄에 붙어 있었던 것이다.",
        before=6,
    )
    add_body(
        doc,
        "전에 “매년 8월 18일”이었을 때는 “8월”이라는 글자가 이 모순을 가리고 있었다. "
        "D-day가 그걸 벗겨냈다. 표시를 바꾼 것이 문제를 만든 게 아니라 원래 있던 문제를 드러냈다.",
        before=6,
    )
    add_body(
        doc,
        "기획서와 대조해 보니 계산은 기획대로였다. 매달 구독은 날짜를 안 보고 전부 더하고, "
        "매년 구독은 결제월이 이번 달일 때만 더한다. 이렇게 정한 이유도 분명하다. "
        "날짜를 보고 판단하면 매달 구독이 결제 다음 날부터 총액에서 빠져, 말일엔 총액이 0원이 된다. "
        "틀린 건 계산이 아니라 “이번 달”이라는 말이 무엇을 가리키는지가 정해지지 않았다는 점이다.",
        before=6,
    )

    add_heading_styled(doc, "7. 위젯을 임시로 막았다", 2)
    add_body(
        doc,
        "“이번 달”을 제대로 정의하기 전에는 위젯을 내보내지 않기로 했다. "
        "위젯이 이상해 보인 건 위젯 잘못이 아니라 이 정의가 흔들려서다. 정의가 정해지면 같이 풀린다.",
    )
    add_body(
        doc,
        "코드는 한 줄도 지우지 않았다. 앱에 위젯을 끼워 넣는 단계에서만 뺐다. "
        "SubscriptionWidget 폴더도, App Group 설정도, entitlements도 그대로 있다.",
        before=6,
    )

    add_heading_styled(doc, "7.1 되살리는 방법", 3)
    add_body(
        doc,
        "SubscriptionTracker.xcodeproj/project.pbxproj를 열고 "
        "Embed Foundation Extensions 단계의 files 괄호 안에 아래 한 줄을 되돌려 넣으면 된다.",
    )
    make_table(
        doc,
        ["항목", "내용"],
        [
            ["찾을 곳", "PBXCopyFilesBuildPhase 안의 “Embed Foundation Extensions”"],
            ["지금 상태", "files = ( ); — 비어 있음"],
            [
                "되돌릴 줄",
                "EF50D92D30348E9A00267E2F /* SubscriptionWidgetExtension.appex "
                "in Embed Foundation Extensions */,",
            ],
            ["확인 방법", "빌드한 .app 안에 PlugIns 폴더가 다시 생기면 성공"],
            ["git으로 하려면", "git diff에 이 한 줄만 찍혀 있으므로 해당 변경만 되돌리면 된다"],
        ],
        [3.6, 13.0],
    )
    add_body(
        doc,
        "Xcode에서 하려면 SubscriptionTracker 타깃 → Build Phases → Embed Foundation Extensions에 "
        "SubscriptionWidgetExtension을 다시 넣으면 같은 결과가 된다.",
        before=8,
    )
    add_body(
        doc,
        "기능을 막을 때 코드를 지우지 않은 이유가 있다. 지우면 되살릴 때 다시 만들어야 하고, "
        "그 사이에 왜 그렇게 짰는지를 잊는다. 끼우는 단계에서만 빼면 되돌리는 비용이 한 줄이다.",
        before=6,
    )

    add_heading_styled(doc, "8. 기획서를 고치다 발견한 것", 2)
    add_body(
        doc,
        "D-day 규칙을 v1.1.0 기획서에 넣으려고 문서를 열었더니, 문서와 구현이 여러 군데 어긋나 있었다. "
        "구현하면서 바꾼 것을 문서에 반영하지 않은 자리다.",
    )
    make_table(
        doc,
        ["항목", "문서", "실제"],
        [
            [
                "위젯 번들 ID",
                "com.dm.SubscriptionTracker.Widget",
                "com.dm.SubscriptionTracker.SubscriptionWidget",
            ],
            ["총액 글자 크기", "작은 22pt / 중간 26pt", "작은 28pt / 중간 22pt (반대)"],
            ["작은 위젯 내용", "다음 결제 이름·날짜 표시", "표시하지 않음"],
            ["빈 상태", "한 가지", "두 가지 (구독 없음 / 이번 달 결제 없음)"],
            ["디자인 토큰", "5개", "10개. 에셋 카탈로그가 아니라 코드로 만듦"],
        ],
        [3.6, 6.6, 6.4],
    )
    add_body(
        doc,
        "D-day는 앱과 위젯 양쪽에 걸리는 규칙이라 위젯 사양 안에 끼우지 않고 4장을 새로 만들었다. "
        "“매년”을 붙였다 뺀 것도 지우지 않고 기록으로 남겼다. "
        "이유까지 적어 두면 나중에 “월/연 구분이 없네”라는 말이 나와도 처음부터 논의하지 않아도 된다.",
        before=8,
    )

    add_heading_styled(doc, "9. 오늘 배운 것", 2)
    add_number(
        doc,
        "표시 형식을 바꾸면 숨어 있던 정의 문제가 드러난다. "
        "“매년 8월 18일”이 “D-364”가 되자 총액과 시제가 안 맞는다는 게 보였다. "
        "화면이 이상하면 대개 화면 탓이 아니다.",
    )
    add_number(
        doc,
        "두 숫자를 나란히 보여줄 거면 같은 기준으로 계산해야 한다. "
        "총액은 이번 달, 개수는 전체였던 것처럼, 각자 맞아도 붙여 놓으면 틀린 말이 된다.",
    )
    add_number(
        doc,
        "고치면 다음 문제가 따라 나온다. 개수 기준을 맞추자 빈 상태 문구가 틀리게 됐다. "
        "한 군데를 바꾸면 그 값을 쓰는 곳을 전부 되짚어야 한다.",
    )
    add_number(
        doc,
        "공용 코드를 한곳에 모아 둔 값이 오늘 나왔다. DateFormat 한 군데를 고치니 앱과 위젯이 같이 바뀌었다.",
    )
    add_number(
        doc,
        "기능을 잠글 때는 지우지 말고 연결을 끊는다. 되돌리는 비용이 한 줄이면 마음 편히 잠글 수 있다.",
    )
    add_number(
        doc,
        "문서는 쓰고 나면 그날부터 어긋나기 시작한다. 구현하면서 바꾼 것은 그때 문서에 적어야 한다. "
        "오늘처럼 몰아서 대조하면 다섯 군데가 한꺼번에 나온다.",
    )

    add_heading_styled(doc, "10. 남은 것 / 다음에 할 일", 2)
    make_table(
        doc,
        ["순서", "할 일", "메모"],
        [
            [
                "1",
                "“이번 달” 큰 숫자의 정의를 정한다",
                "후보: 이번 달 고정지출 유지 / 앞으로 30일 / 월 평균 환산",
            ],
            ["2", "정의가 정해지면 위젯을 되살린다", "pbxproj 한 줄 (7.1 참고)"],
            ["3", "v1.1.0 기획서의 시안 이미지를 다시 그린다", "지금 그림은 옛 안이라 실제와 다름"],
            ["4", "MARKETING_VERSION을 1.0에서 1.1.0으로 올린다", "앱·위젯 두 타깃 모두"],
            ["5", "Apple Developer Program 승인 확인", "신분증 제출 후 대기 중"],
        ],
        [1.6, 7.4, 7.6],
    )

    add_heading_styled(doc, "11. 오늘 만진 파일", 2)
    make_table(
        doc,
        ["파일", "무엇을"],
        [
            ["Shared/Formatters.swift", "D-day 포맷 추가. 날짜 포맷은 남겨 둠"],
            ["Shared/MonthlyTotal.swift", "이번 달 포함 여부를 includes로 분리"],
            ["Shared/DesignTokens.swift", "위젯 색·글자 토큰 (이번 세션 이전 작업분 포함)"],
            ["SubscriptionWidget/SubscriptionWidget.swift", "대표 이름, 개수 기준, 빈 상태 두 가지, D-day"],
            ["SubscriptionTracker/SubscriptionListView.swift", "행 날짜를 D-day로"],
            ["SubscriptionTracker/AddSubscriptionView.swift", "저장 시 결제일 밀어내기"],
            ["SubscriptionTracker/SubscriptionDetailView.swift", "저장 시 결제일 밀어내기"],
            ["SubscriptionTracker.xcodeproj/project.pbxproj", "위젯 임베드 제외 (임시)"],
            ["pyfile/build_v110_planning.py", "4장 신설, 문서·구현 불일치 5건 수정"],
        ],
        [7.6, 9.0],
    )

    add_heading_styled(doc, "12. 확인에 쓴 방법", 2)
    add_bullet(doc, "xcodebuild -project … -scheme … -destination 'platform=iOS Simulator,name=iPhone 17' build")
    add_bullet(doc, "xcrun simctl install booted <경로>/SubscriptionTracker.app")
    add_bullet(doc, "xcrun simctl io booted screenshot 파일.png")
    add_bullet(doc, "xcrun simctl get_app_container booted <번들ID> groups — App Group 폴더 찾기")
    add_bullet(doc, "sqlite3 default.store “select …” — 화면 숫자가 맞는지 저장값과 대조")
    add_body(
        doc,
        "총액 188,000원이 맞는지 확인할 때 sqlite3로 저장된 다섯 건을 직접 읽어 더해 봤다. "
        "화면만 보고 넘어가면 계산이 맞는지 알 수 없다.",
        before=6,
    )

    os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
    doc.save(OUT_PATH)


if __name__ == "__main__":
    build()
    print(OUT_PATH)
