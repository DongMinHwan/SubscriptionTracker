#!/usr/bin/env python3
"""2026-08-18 개발일지 .docx 생성."""

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

from build_docx import (
    add_body,
    add_bullet,
    add_heading_styled,
    add_number,
    configure_document,
    make_table,
    set_paragraph_format,
    style_run,
)

BASE = "/Users/hwangdongmin/Projects/NewProject_1"
OUT_PATH = f"{BASE}/개발일지/2026-08-18_개발일지.docx"
IMAGE_DIR = f"{BASE}/개발일지/이미지"

KV = [3.2, 13.4]


def add_picture(doc, filename, width_cm, caption):
    p = doc.add_paragraph()
    set_paragraph_format(p, before=8, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
    p.add_run().add_picture(os.path.join(IMAGE_DIR, filename), width=Cm(width_cm))

    cap = doc.add_paragraph()
    set_paragraph_format(cap, before=0, after=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    style_run(cap.add_run(caption), 10, False)


def add_picture_pair(doc, left, right, width_cm=5.0):
    """스크린샷 두 장을 나란히 놓고 아래에 설명을 단다."""
    table = doc.add_table(rows=2, cols=2)
    table.autofit = True

    for col, (filename, caption) in enumerate((left, right)):
        cell = table.rows[0].cells[col]
        p = cell.paragraphs[0]
        set_paragraph_format(p, before=4, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
        p.add_run().add_picture(os.path.join(IMAGE_DIR, filename), width=Cm(width_cm))

        cap = table.rows[1].cells[col].paragraphs[0]
        set_paragraph_format(cap, before=0, after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        style_run(cap.add_run(caption), 10, False)


def build():
    doc = Document()
    configure_document(doc)

    add_heading_styled(doc, "개발일지 — 2026년 8월 18일 (화)", 1)

    make_table(
        doc,
        ["항목", "내용"],
        [
            ["날짜", "2026년 8월 18일 화요일"],
            ["프로젝트", "구독 트래커 (학습용 첫 iOS 앱)"],
            ["오늘 목표", "출시를 막는 것들을 치운다."],
            [
                "결과",
                "앱 이름·글자 크기·다크 모드 정리. Apple Developer Program 결제 완료(승인 대기).",
            ],
            [
                "한 줄 소감",
                "원인이라고 생각했던 곳이 원인이 아니었다. 확인해보길 잘했다.",
            ],
        ],
        KV,
    )

    add_heading_styled(doc, "1. 지난 일지 이후 (8월 17일 저녁)", 2)
    add_body(
        doc,
        "어제 일지를 쓴 뒤에도 작업이 이어져서, 8월 17일 일지에 빠진 것을 먼저 적어 둔다.",
    )
    make_table(
        doc,
        ["시각", "한 일"],
        [
            [
                "8/17 16:22",
                "배포 타겟을 iOS 26.2에서 18.0으로 낮추고 iPad 지원을 껐다. "
                "26.2로 두면 설치 가능한 기기가 거의 없다. 코드 수정 없이 그대로 빌드됐다.",
            ],
            [
                "8/17 16:48",
                "git 저장소를 만들었다(git init).",
            ],
            [
                "8/17 17:08",
                ".gitignore를 쓰고 첫 커밋(a20301b). 37개 파일.",
            ],
            [
                "8/17 17:14",
                "GitHub에 비공개 저장소 SubscriptionTracker를 만들고 push. 56개 객체, 1.88 MiB.",
            ],
        ],
        [2.6, 14.0],
    )

    add_heading_styled(doc, "2. 오늘의 흐름", 2)
    make_table(
        doc,
        ["시각", "한 일", "산출물"],
        [
            [
                "오전",
                "Apple Developer Program 등록 및 결제 (129,000원)",
                "승인 대기",
            ],
            ["21:11", "가입 상태 확인. Pending", "-"],
            [
                "21:15 ~ 21:24",
                "출시 전 정리 3가지: 앱 표시 이름, 글자 크기 대응, 다크 모드 판단",
                "파일 4개 수정",
            ],
            ["21:26", "커밋 f0c0b3c", "커밋 2개째"],
        ],
        [2.6, 9.4, 4.6],
    )

    add_heading_styled(doc, "3. Apple Developer Program", 2)
    make_table(
        doc,
        ["항목", "내용"],
        [
            ["비용", "연 129,000원. 1년마다 갱신. 놓치면 앱이 스토어에서 내려간다."],
            ["계정", "회사 팀에서 쓰던 Apple 계정에 개인 멤버십을 추가하는 방식"],
            ["상태", "결제 완료, 승인 대기(Pending). 보통 24~48시간"],
        ],
        KV,
    )
    add_body(
        doc,
        "헷갈렸던 것 하나. App Store Connect에 이미 회사 팀 세 곳으로 들어가 있어서 "
        "출시할 수 있는 줄 알았는데, 회사 팀의 App Manager 권한과 내 앱을 낼 권한은 별개였다. "
        "개인 멤버십은 따로 가입해야 한다.",
        before=8,
    )
    add_body(
        doc,
        "결제 후 화면에 “Purchase your membership”이 계속 떠서 결제가 실패한 줄 알았다. "
        "링크를 바로 누르면 중복 결제될 수 있어서, 카드 승인 내역부터 확인하고 기다리기로 했다.",
        before=6,
    )

    add_heading_styled(doc, "4. 앱 표시 이름", 2)
    add_body(
        doc,
        "홈 화면에서 이름이 SubscriptionTra… 로 잘리고 있었다. "
        "Info.plist에 CFBundleDisplayName을 “구독 트래커”로 넣어 해결했다.",
    )
    add_body(
        doc,
        "제출할 때마다 뜨는 암호화 질문을 막으려고 ITSAppUsesNonExemptEncryption도 false로 함께 넣었다.",
        before=6,
    )
    add_picture(doc, "2026-08-18_앱이름.png", 4.2, "홈 화면 (수정 후)")

    add_heading_styled(doc, "5. 글자 크기 — 원인이 다른 곳에 있었다", 2)
    add_body(
        doc,
        "앱 진입점에 .dynamicTypeSize(.large)를 걸어 둬서 사용자의 글자 크기 설정이 "
        "무시되는 줄 알았다. 그래서 그 줄을 지우고 시뮬레이터에서 최대 접근성 크기로 바꿔 봤다.",
    )
    add_body(doc, "그런데 글자가 하나도 안 커졌다. 오른쪽 화살표만 커졌다.", before=6)
    add_body(
        doc,
        "진짜 원인은 DesignTokens였다. Font.system(size: 17)은 고정 크기 폰트라 "
        "Dynamic Type을 아예 무시한다. 잠금을 푸는 것만으로는 아무 일도 일어나지 않았던 것이다. "
        "화살표만 커진 건 그건 시스템이 그리는 요소이기 때문이다.",
        before=6,
    )

    add_heading_styled(doc, "5.1 고친 방법", 3)
    add_body(
        doc,
        "기본 크기가 기획서 pt 값과 똑같은 텍스트 스타일로 바꿨다. "
        "덕분에 기본 상태의 모양은 그대로 두면서 확대만 따라가게 됐다.",
    )
    make_table(
        doc,
        ["토큰", "기존", "변경", "기본 크기"],
        [
            ["totalLabel", "system(size: 13)", ".footnote", "13pt"],
            ["totalValue", "system(size: 34, bold)", ".largeTitle.bold()", "34pt"],
            ["rowTitle / rowAmount", "system(size: 17)", ".body", "17pt"],
            ["rowSubtitle", "system(size: 13)", ".footnote", "13pt"],
            ["widgetLabel", "system(size: 12)", ".caption", "12pt"],
            ["widgetValue", "system(size: 22, bold)", ".title2.bold()", "22pt"],
        ],
        [4.2, 4.6, 4.0, 3.8],
    )
    add_body(
        doc,
        "행 높이도 .frame(height: 60)에서 minHeight로 바꿨다. 고정이면 글자가 커져도 행이 안 늘어난다.",
        before=8,
    )

    add_heading_styled(doc, "5.2 그랬더니 다음 문제가 나왔다", 3)
    add_body(
        doc,
        "최대 크기에서 금액이 ₩17,0 / 00 으로 쪼개졌다. 돈이 중간에서 잘리는 건 안 된다. "
        "글자가 아주 커지면(dynamicTypeSize.isAccessibilitySize) 가로 배치를 포기하고 "
        "이름·날짜·금액을 세로로 쌓게 했다.",
    )
    add_picture_pair(
        doc,
        ("2026-08-18_수정전_최대글자.png", "수정 전: 글자가 안 커지고 금액이 잘림"),
        ("2026-08-18_수정후_최대글자.png", "수정 후: 세로로 쌓아 금액이 온전함"),
    )
    add_picture(doc, "2026-08-18_큰글자.png", 5.0, "중간 크기(XXXL). 가로 배치 유지")

    add_heading_styled(doc, "6. 다크 모드 — 유지하기로 했다", 2)
    add_body(
        doc,
        "다크 모드 기기에서 실행해 보니 깨지는 데 없이 라이트로 잘 나왔다. "
        ".preferredColorScheme(.light)를 그대로 두기로 했다.",
    )
    add_body(
        doc,
        "색 토큰 7개가 전부 라이트 값으로 고정돼 있어서, 잠금을 풀면 오히려 문제가 생긴다. "
        "내용은 밝은데 내비게이션 바만 어두워지는 어색한 조합이 된다. "
        "제대로 하려면 토큰마다 다크 값을 새로 정해야 하는데, 그건 v1.1 몫이다.",
        before=6,
    )
    add_body(
        doc,
        "v1.0.1 기획서에도 “다크 모드 없음. 라이트 고정”이라고 적혀 있으니 기획대로 간다. "
        "빼먹은 게 아니라 정한 것이라는 점이 중요하다.",
        before=6,
    )

    add_heading_styled(doc, "7. 오늘 배운 것", 2)
    add_number(
        doc,
        "SwiftUI의 Font.system(size:)는 고정 크기다. Dynamic Type에 대응하려면 "
        ".body, .footnote 같은 텍스트 스타일을 써야 한다.",
    )
    add_number(
        doc,
        "원인이라고 생각한 곳을 고치기 전에 실제로 재현해 보는 게 빠르다. "
        "오늘은 진입점 한 줄을 지워도 아무 변화가 없어서 진짜 원인을 찾을 수 있었다.",
    )
    add_number(
        doc,
        "접근성 대응은 글자를 키우는 것으로 끝나지 않는다. 커진 다음의 배치까지 정해야 한다.",
    )
    add_number(
        doc,
        "회사 팀 소속과 개인 개발자 멤버십은 별개다. App Store Connect에 들어가진다고 "
        "내 앱을 낼 수 있는 게 아니다.",
    )
    add_number(
        doc,
        "git init과 commit은 전부 로컬이다. remote를 등록하고 push해야 밖으로 나간다.",
    )

    add_heading_styled(doc, "8. 확인에 쓴 방법", 2)
    add_body(doc, "화면을 눈으로 확인하려고 쓴 명령들. 다음에 또 쓸 것 같아 적어 둔다.", after=4)
    add_bullet(doc, "xcrun simctl ui <기기> content_size large | accessibility-extra-extra-extra-large")
    add_bullet(doc, "xcrun simctl ui <기기> appearance light | dark")
    add_bullet(doc, "xcrun simctl io <기기> screenshot 파일.png")
    add_body(
        doc,
        "목록이 비어 있으면 확인이 안 돼서, 시뮬레이터의 SwiftData 저장소에 sqlite3로 "
        "구독 4개를 직접 넣고 확인했다. 실기기와는 무관하다.",
        before=6,
    )

    add_heading_styled(doc, "9. 남은 것 / 다음에 할 일", 2)
    make_table(
        doc,
        ["순서", "할 일", "메모"],
        [
            ["1", "Apple Developer Program 승인 기다리기", "48시간 넘으면 문의"],
            [
                "2",
                "개인정보 처리방침 URL 만들기",
                "데이터를 수집하지 않아도 필수. GitHub Pages면 충분",
            ],
            ["3", "App Store Connect에 앱 등록", "이름, 설명, 키워드, 카테고리"],
            [
                "4",
                "스토어 스크린샷 만들기",
                "6.9인치 1320×2868 한 세트. 알파 채널 금지",
            ],
            ["5", "Archive 후 업로드, 심사 제출", "심사는 90%가 24시간 이내"],
        ],
        [1.6, 7.0, 8.0],
    )
    add_body(
        doc,
        "스크린샷을 고르다 보면 어색한 화면이 저절로 보일 것이다. "
        "완성도 목록은 그때 만들기로 한다. 지금 상상으로 만드는 것보다 정확하다.",
        before=8,
    )

    add_heading_styled(doc, "10. 오늘 만진 파일", 2)
    make_table(
        doc,
        ["파일", "무엇을"],
        [
            ["SubscriptionTracker/Info.plist", "앱 표시 이름, 암호화 선언 추가"],
            ["DesignTokens.swift", "폰트 토큰을 텍스트 스타일로 교체"],
            [
                "SubscriptionListView.swift",
                "행 높이 minHeight로, 큰 글자에서 세로 배치",
            ],
            ["SubscriptionTrackerApp.swift", "dynamicTypeSize 고정 제거"],
        ],
        [7.0, 9.6],
    )
    add_body(doc, "커밋: f0c0b3c “출시 준비 : 앱이름 표시, 글자크기대응, 암호화 선언”", before=8)

    os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
    doc.save(OUT_PATH)


if __name__ == "__main__":
    build()
    print(OUT_PATH)
