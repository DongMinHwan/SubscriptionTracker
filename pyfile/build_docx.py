#!/usr/bin/env python3
"""Build a formatted .docx from the subscription-tracker research notes."""

from urllib.parse import quote, urlsplit, urlunsplit

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.opc.constants import RELATIONSHIP_TYPE as RT

OUT_PATH = "/Users/hwangdongmin/Projects/NewProject_1/구독_트래커_유사앱_조사.docx"
TABLE_WIDTH_CM = 16.6
HEADER_FILL = "1F4E79"
HEADER_RGB = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BODY_RGB = RGBColor(0x22, 0x22, 0x22)
TWIPS_PER_CM = 567


def east_asia_font(run, name_en="Malgun Gothic", name_kr="맑은 고딕"):
    run.font.name = name_en
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), name_en)
    rFonts.set(qn("w:hAnsi"), name_en)
    rFonts.set(qn("w:eastAsia"), name_kr)
    rFonts.set(qn("w:cs"), name_en)


def style_run(run, size=11, bold=False, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color if color is not None else BODY_RGB
    east_asia_font(run)


def set_paragraph_format(paragraph, *, before=0, after=8, line=1.15, align=None):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.widow_control = True
    if align is not None:
        paragraph.alignment = align


def add_body(doc, text, *, before=0, after=8):
    p = doc.add_paragraph()
    set_paragraph_format(p, before=before, after=after)
    run = p.add_run(text)
    style_run(run, 11, False)
    return p


def add_heading_styled(doc, text, level):
    sizes = {1: 18, 2: 14, 3: 12}
    spacing = {
        1: (0, 12),
        2: (16, 8),
        3: (12, 6),
    }
    p = doc.add_paragraph()
    p.style = doc.styles[f"Heading {level}"]
    before, after = spacing[level]
    set_paragraph_format(p, before=before, after=after, line=1.15)
    # clear default runs
    if p.runs:
        p.runs[0].text = ""
    run = p.add_run(text)
    style_run(run, sizes[level], True, HEADER_RGB)
    return p


def _empty_paragraph(p):
    for child in list(p._element):
        if child.tag == qn("w:r"):
            p._element.remove(child)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    _empty_paragraph(p)
    set_paragraph_format(p, before=0, after=4, line=1.15)
    p.paragraph_format.left_indent = Cm(0.75)
    run = p.add_run(text)
    style_run(run, 11, False)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    _empty_paragraph(p)
    set_paragraph_format(p, before=0, after=4, line=1.15)
    p.paragraph_format.left_indent = Cm(0.75)
    run = p.add_run(text)
    style_run(run, 11, False)
    return p


def set_style_font(style, size, bold=False):
    style.font.name = "Malgun Gothic"
    style.font.size = Pt(size)
    style.font.bold = bold
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), "Malgun Gothic")
    rFonts.set(qn("w:hAnsi"), "Malgun Gothic")
    rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    rFonts.set(qn("w:cs"), "Malgun Gothic")


def shade_cell(cell, fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag == qn("w:shd"):
            tcPr.remove(child)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def set_cell_margins(cell, top=80, bottom=80, left=80, right=80):
    """Margins in twips. 4pt = 80 twips."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag == qn("w:tcMar"):
            tcPr.remove(child)
    tcMar = OxmlElement("w:tcMar")
    for name, value in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{name}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def set_cell_width(cell, width_cm):
    twips = int(round(width_cm * TWIPS_PER_CM))
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag == qn("w:tcW"):
            tcPr.remove(child)
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(twips))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)
    cell.width = Cm(width_cm)


def set_table_fixed_width(table, widths_cm):
    widths_cm = list(widths_cm)
    diff = round(TABLE_WIDTH_CM - sum(widths_cm), 10)
    if abs(diff) >= 0.001:
        widths_cm[-1] = round(widths_cm[-1] + diff, 10)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    total_twips = int(round(sum(widths_cm) * TWIPS_PER_CM))
    tblW.set(qn("w:w"), str(total_twips))
    tblW.set(qn("w:type"), "dxa")

    jc = tblPr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        tblPr.append(jc)
    jc.set(qn("w:val"), "center")

    layout = tblPr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.find(qn("w:tblGrid"))
    if grid is not None:
        tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    for w in widths_cm:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(round(w * TWIPS_PER_CM))))
        grid.append(col)
    tbl.insert(1, grid)

    for row in table.rows:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        cant = OxmlElement("w:cantSplit")
        trPr.append(cant)
        for i, w in enumerate(widths_cm):
            set_cell_width(row.cells[i], w)


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), HEADER_FILL)
        borders.append(el)
    tblPr.append(borders)


def clear_cell(cell):
    for p in cell.paragraphs:
        p.clear()
    while len(cell.paragraphs) > 1:
        el = cell.paragraphs[-1]._element
        el.getparent().remove(el)


def write_cell(cell, text, *, header=False, align="left", bold=False):
    clear_cell(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
    lines = text.split("\n") if text else [""]
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        set_paragraph_format(
            p,
            before=0,
            after=0 if i == len(lines) - 1 else 2,
            line=1.15,
            align=(
                WD_ALIGN_PARAGRAPH.CENTER
                if header or align == "center"
                else WD_ALIGN_PARAGRAPH.LEFT
            ),
        )
        run = p.add_run(line)
        if header:
            style_run(run, 11, True, WHITE)
        else:
            style_run(run, 11, bold, BODY_RGB)
    if header:
        shade_cell(cell, HEADER_FILL)
    else:
        shade_cell(cell, "FFFFFF")


def encode_url(url):
    parts = urlsplit(url)
    path = quote(parts.path, safe="/-._~")
    return urlunsplit((parts.scheme, parts.netloc, path, parts.query, parts.fragment))


def add_hyperlink(paragraph, text, url):
    encoded = encode_url(url)
    part = paragraph.part
    r_id = part.relate_to(encoded, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "22")
    rPr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), "22")
    rPr.append(szCs)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Malgun Gothic")
    rFonts.set(qn("w:hAnsi"), "Malgun Gothic")
    rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    rFonts.set(qn("w:cs"), "Malgun Gothic")
    rPr.append(rFonts)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def make_table(doc, headers, rows, widths_cm, *, header_align="center"):
    assert len(headers) == len(widths_cm)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    set_table_borders(table)
    set_table_fixed_width(table, widths_cm)
    for i, h in enumerate(headers):
        write_cell(table.rows[0].cells[i], h, header=True, align=header_align)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            first_col_kv = len(headers) == 2 and headers[0] == "항목"
            write_cell(
                table.rows[r_idx + 1].cells[c_idx],
                val,
                header=False,
                align="left",
                bold=(first_col_kv and c_idx == 0),
            )
    tr = table.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)
    return table


def make_link_table(doc, items, widths_cm):
    table = doc.add_table(rows=1 + len(items), cols=2)
    set_table_borders(table)
    set_table_fixed_width(table, widths_cm)
    write_cell(table.rows[0].cells[0], "이름", header=True)
    write_cell(table.rows[0].cells[1], "URL", header=True)
    tr = table.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)
    for i, (name, url) in enumerate(items):
        write_cell(table.rows[i + 1].cells[0], name, header=False, bold=True)
        cell = table.rows[i + 1].cells[1]
        clear_cell(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        shade_cell(cell, "FFFFFF")
        p = cell.paragraphs[0]
        set_paragraph_format(p, before=0, after=0, line=1.15)
        add_hyperlink(p, url, url)
    return table


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    normal = doc.styles["Normal"]
    set_style_font(normal, 11, False)
    normal.font.color.rgb = BODY_RGB
    pf = normal.paragraph_format
    pf.line_spacing = 1.15
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.space_after = Pt(8)
    pf.space_before = Pt(0)

    for level, size in ((1, 18), (2, 14), (3, 12)):
        st = doc.styles[f"Heading {level}"]
        set_style_font(st, size, True)
        st.font.color.rgb = HEADER_RGB
        st.paragraph_format.line_spacing = 1.15
        st.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

    # Default theme east-Asia font
    styles = doc.styles.element
    rpr_default = styles.find(qn("w:docDefaults"))
    if rpr_default is not None:
        rpr = rpr_default.find(qn("w:rPrDefault"))
        if rpr is not None:
            rPr = rpr.find(qn("w:rPr"))
            if rPr is None:
                rPr = OxmlElement("w:rPr")
                rpr.append(rPr)
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.append(rFonts)
            rFonts.set(qn("w:ascii"), "Malgun Gothic")
            rFonts.set(qn("w:hAnsi"), "Malgun Gothic")
            rFonts.set(qn("w:eastAsia"), "맑은 고딕")
            rFonts.set(qn("w:cs"), "Malgun Gothic")


def build():
    doc = Document()
    configure_document(doc)

    add_heading_styled(doc, "구독·고정지출 트래커 유사 앱 조사", 1)

    kv2 = [3.2, 13.4]
    make_table(
        doc,
        ["항목", "내용"],
        [
            ["작성일", "2026년 8월 15일"],
            [
                "목적",
                "iOS 네이티브로 구독·고정지출 트래커를 만들 때, 이미 있는 유사 앱을 정리하고 차별 지점을 찾기 위함.",
            ],
            [
                "조사 범위",
                "미국·한국 App Store, 앱 공식 사이트, 2026년 비교 리뷰, 한국 핀테크(토스·뱅크샐러드·카카오페이 등) 기능.",
            ],
        ],
        kv2,
    )

    add_heading_styled(doc, "1. 한 줄 결론", 2)
    add_body(
        doc,
        "비슷한 앱이 이미 많다. 특히 “수동 입력 + 이번 달 합계 + 결제 전 알림 + 홈 화면 위젯”은 2026년 iOS 인디 앱에서 가장 흔한 구독 트래커 형태다.",
    )
    add_body(
        doc,
        "연습용으로 기획–디자인–개발–출시–운영을 끝까지 가려면 그대로 해도 된다. 스토어에서 이기려면 각도를 바꿔야 한다. 한국 사용자는 보통 토스·뱅크샐러드·카카오페이처럼 카드에서 숨은 결제를 찾아주는 쪽을 먼저 쓴다.",
    )

    add_heading_styled(doc, "2. 당초 추천안 (비교 기준)", 2)
    add_body(
        doc,
        "혼자 iOS 개발자가 기획부터 운영까지 가려면, 너무 작지도 않고 백엔드가 필수여서 막히지도 않는 제품이 필요하다는 전제로 아래를 추천했다.",
    )
    add_bullet(doc, "스택: iOS 네이티브 (SwiftUI). Android·크로스는 사용자가 생긴 뒤 재검토.")
    add_bullet(doc, "제품: 서버 없는 구독·고정지출 트래커.")
    add_bullet(
        doc,
        "v1 기능 3개: 구독 추가/수정(이름, 금액, 주기, 다음 결제일) / 이번 달 총액과 다가오는 결제 / 결제 전날 알림 + 홈 화면 위젯.",
    )
    add_bullet(
        doc,
        "v1에서 빼기로 한 것: 가계부, 은행 연동, 소셜 로그인, 클라우드 동기화, Android, 복잡한 차트.",
    )
    add_body(
        doc,
        "한 문장 정의: 매달 나가는 구독을 한곳에 모아, 이번 달 얼마인지 홈 화면에서 바로 보게 한다.",
        before=6,
    )

    add_heading_styled(doc, "3. 시장을 나누는 두 갈래", 2)
    make_table(
        doc,
        ["유형", "하는 일", "대표 앱", "추천안과의 관계"],
        [
            [
                "수동 입력 / 개인정보 우선",
                "내가 아는 구독을 직접 넣고 월 합계·알림·위젯으로 관리",
                "Bobby\nSubo\nSubTrack'd\nSubManager\nSubvise\nSubby",
                "기능이 거의 동일. 가장 직접적인 경쟁.",
            ],
            [
                "자동 발견 / 금융 연동",
                "계좌·카드 내역에서 숨은 정기결제를 찾아줌",
                "토스\n뱅크샐러드\n카카오페이\n왓섭\nRocket Money",
                "다른 제품. 한국에서는 이쪽이 더 많이 쓰임.",
            ],
        ],
        [3.2, 5.0, 4.0, 4.3],
    )
    add_body(
        doc,
        "핵심 차이: 인디 앱은 “정리”를 잘하고, 핀테크는 “발견”을 잘한다. 한국 사용자는 발견 문제를 먼저 느낀다.",
        before=8,
    )

    add_heading_styled(doc, "4. 거의 같은 제품 (수동 입력, 은행 연동 없음)", 2)
    add_body(
        doc,
        "추천안과 기능이 가장 겹친다. 계정 없이 직접 넣고, 이번 달 얼마인지 보고, 갱신 전에 알림을 받는다.",
    )

    add_heading_styled(doc, "4.1 Bobby", 3)
    make_table(
        doc,
        ["항목", "내용"],
        [
            ["플랫폼", "iOS"],
            ["핵심", "미니멀 리스트, 커스텀 아이콘, 월 합계. 이 카테고리의 기준점."],
            ["가격", "무료 + 일회성 Pro(올인원 팩)"],
            ["평가", "미국 App Store 기준 약 8,000개 리뷰, 4.7점대"],
            ["만든 곳", "암스테르담 디자인 스튜디오 Yummygum"],
            [
                "특징",
                "은행 연결 없음. CSV/AI 가져오기·해지 가이드는 없음. 단순함을 의도적으로 지킴.",
            ],
            ["겹치는 정도", "거의 동일"],
        ],
        kv2,
    )
    add_body(
        doc,
        "참고: 일부 리뷰는 업데이트가 느려졌고, 알림·분석이 약하다고 본다. 그래도 “예쁜 구독 리스트”의 대표작이다.",
        before=8,
    )

    add_heading_styled(doc, "4.2 Subo", 3)
    make_table(
        doc,
        ["항목", "내용"],
        [
            ["플랫폼", "iOS (한국 앱스토어 등록됨)"],
            ["한국 스토어명", "구독 관리·갱신 알림 트래커: Subo"],
            ["핵심", "갱신 알림, 홈/잠금 위젯, 인기 서비스 템플릿, 다중 통화, iCloud"],
            [
                "가격(한국 스토어 표시 기준)",
                "Unlimited 평생권 약 29,000원 / 3개월 약 6,600원 / 12개월 약 15,000원",
            ],
            [
                "개인정보",
                "기기 저장, 선택 시 iCloud. 광고·데이터 판매 없음. 독립 개발.",
            ],
            [
                "겹치는 정도",
                "거의 동일. 한국 사용자가 바로 검색해 찾을 수 있는 직접 경쟁작.",
            ],
        ],
        kv2,
    )

    add_heading_styled(doc, "4.3 SubTrack'd", 3)
    make_table(
        doc,
        ["항목", "내용"],
        [
            ["플랫폼", "iOS (iOS 26 대응을 내세움)"],
            [
                "핵심",
                "홈·잠금·StandBy 위젯, Siri, 스크린샷 OCR, 지출 인사이트, 선택적 iCloud",
            ],
            ["가격", "무료(구독 5개) + Pro 일회성 $5.99"],
            ["포지션", "Bobby가 최신 iOS에 덜 맞춰진 틈을 노린 Bobby 대안"],
            ["겹치는 정도", "거의 동일 + 위젯·Siri를 더 앞세움"],
        ],
        kv2,
    )

    add_heading_styled(doc, "4.4 SubManager", 3)
    make_table(
        doc,
        ["항목", "내용"],
        [
            ["플랫폼", "iPhone, iPad, Mac, Apple Watch, Apple Vision, 웹(iCloud)"],
            ["핵심", "Apple 생태계 동기화, 위젯, Siri Shortcuts, 지출 인사이트"],
            ["가격", "일회성 SubManager+ (약 $7.99, 가족 공유 플랜 있음)"],
            ["특징", "9to5Mac Indie App Spotlight 소개. 계정·광고 없음."],
            ["겹치는 정도", "높음. 멀티 디바이스까지 가면 추천안 v1보다 넓음."],
        ],
        kv2,
    )

    add_heading_styled(doc, "4.5 Subvise, Subby", 3)
    make_table(
        doc,
        ["항목", "내용"],
        [
            [
                "Subvise",
                "카드형 목록, 100개 이상 서비스 카탈로그, 무료 5개 + 일회성 Pro. 약정·해지 통보 기간 추적.",
            ],
            [
                "Subby",
                "iOS+Android. 수동 입력, 알림. 무료(광고) 또는 저가 일회성/구독형 언급이 자료마다 다름.",
            ],
            ["겹치는 정도", "높음"],
        ],
        kv2,
    )

    add_heading_styled(doc, "4.6 기능이 더 많은 수동형: Subcut, ReSubs, SubTracker", 3)
    make_table(
        doc,
        ["앱", "추가로 하는 일", "가격 경향", "의미"],
        [
            [
                "Subcut",
                "Gmail 영수증, 은행 명세서 CSV/PDF, App Store 구독 가져오기, 해지 가이드, 캘린더 뷰",
                "구독형(약 $1.99/월 등)",
                "추천안 v1보다 기능이 많음",
            ],
            [
                "ReSubs",
                "CSV·Gmail·스크린샷 AI 추출, 체험/일시정지/해지 상태, 해지 가이드, iOS+Android",
                "무료 한도 + 프리미엄",
                "크로스플랫폼 수동 트래커의 대표",
            ],
            [
                "SubTracker",
                "카테고리 차트, 해지 가이드·이메일 템플릿, 스크린샷 가져오기. 한국 스토어 있음",
                "무료 3개 + Pro 월/연 구독",
                "한국 스토어에서 바로 검색됨",
            ],
        ],
        [2.4, 6.2, 3.4, 4.5],
    )
    add_body(
        doc,
        "이 그룹은 “리스트 + 알림”을 넘어 가져오기와 해지까지 한다. 같은 카테고리에서 후발 주자가 기능을 쌓는 방향이다.",
        before=8,
    )

    add_heading_styled(doc, "5. 다른 제품 (자동 발견) — 한국에서 실제 대체재", 2)
    add_body(
        doc,
        "추천안과 기능은 다르지만, 사용자가 “구독 정리하려고” 먼저 여는 앱이다. 한국 시장에서는 이쪽이 더 강한 경쟁이다.",
    )

    add_heading_styled(doc, "5.1 토스", 3)
    make_table(
        doc,
        ["항목", "내용"],
        [
            ["방식", "계좌·카드·토스페이 정기결제/자동결제 조회"],
            ["한국 적합성", "가장 흔함"],
            ["강점", "숨은 결제 발견, 일부 해지, 이미 깔려 있음"],
            [
                "약점",
                "위젯으로 “이번 달 고정비”를 의도적으로 모아 두게 하는 UX는 약함. App Store 구독, 해외 카드, 가족 명의, 현금성 고정비는 빠질 수 있음",
            ],
        ],
        kv2,
    )

    add_heading_styled(doc, "5.2 뱅크샐러드", 3)
    make_table(
        doc,
        ["항목", "내용"],
        [
            ["방식", "마이데이터 기반 지출 분석"],
            ["강점", "자산·소비와 함께 봄"],
            ["약점", "구독 전용 앱이 아니라 가계부 안에 구독이 묻힘"],
        ],
        kv2,
    )

    add_heading_styled(doc, "5.3 카카오페이", 3)
    make_table(
        doc,
        ["항목", "내용"],
        [
            ["방식", "카카오페이 정기결제 조회·해지"],
            ["강점", "카톡 안에서 접근하기 쉬움"],
            ["약점", "카카오페이로 결제한 것만 잘 보임"],
        ],
        kv2,
    )

    add_heading_styled(doc, "5.4 왓섭 (Wassup)", 3)
    make_table(
        doc,
        ["항목", "내용"],
        [
            [
                "방식",
                "결제 내역에서 구독 자동 분류. 한국 전용 구독 관리 앱으로 언급됨",
            ],
            [
                "의미",
                "같은 문제를 자동으로 푸는 로컬 경쟁자. 수동 입력 인디 앱보다 한국 사용자 기대에 가까움",
            ],
        ],
        kv2,
    )

    add_heading_styled(doc, "5.5 Rocket Money (구 Truebill) 및 해외 가계부", 3)
    make_table(
        doc,
        ["항목", "내용"],
        [
            [
                "Rocket Money",
                "은행 연동으로 구독 자동 탐지, 해지·요금 협상 대행. 월 구독. 한국 은행과 잘 안 맞음.",
            ],
            [
                "같은 계열",
                "Copilot Money, Monarch Money, PocketGuard, Quicken Simplifi, YNAB(전체 예산)",
            ],
            [
                "참고",
                "직접 경쟁은 약함. 다만 “자동으로 찾아주는 앱”이 글로벌 표준이라는 점은 참고할 만함.",
            ],
        ],
        kv2,
    )

    add_heading_styled(doc, "6. 겹치는 기능 vs 빈 구멍", 2)
    make_table(
        doc,
        ["이미 있는 것", "한국 기준 빈 구멍"],
        [
            [
                "그대로 내면 차별이 약함\n수동 추가, 월/연 합계, 결제일 알림\n홈 화면·잠금 화면 위젯\n넷플릭스 등 인기 서비스 템플릿\niCloud 동기화, Face ID\n해지 가이드, 영수증·이메일 가져오기 (상위 앱)",
                "한국 기준으로 상대적으로 약한 지점\n티빙, 쿠팡와우, 통신비, 보험, 관리비 같은 한국 고정지출을 한 화면에 모으기\n토스가 잘 못 보는 것: App Store 구독, 해외 카드, 현금, 가족 명의 결제\n가구 단위 “우리집 고정비” (개인 카드 스캔과 다름)\n체험판 종료일만 빠르게 심는 초경량 UX\n한국어 카피, 원 단위, 한국 결제 주기(월/연/통신 약정)",
            ],
        ],
        [8.3, 8.3],
    )

    add_heading_styled(doc, "7. 시사점", 2)
    add_number(
        doc,
        "학습·포트폴리오: Bobby/Subo를 참고작으로 두고 SwiftUI, 위젯, 알림, 심사를 끝까지 겪는 용도로는 여전히 적합하다. 스토어에서 눈에 띄기는 어렵다.",
    )
    add_number(
        doc,
        "제품으로 가려면: 구독만이 아니라 통신비·보험·관리비·쿠팡와우·OTT 결합을 “이번 달 고정비”로 모으는 한국 각도가 낫다. 토스는 카드 내역은 잘 보여도, 한곳에 의도적으로 모아 두고 결제 전에 결정하게 만드는 UX는 약하다.",
    )
    add_number(
        doc,
        "카테고리 변경도 검토: N빵 정산, 유통기한, 체험판만 추적처럼 처음부터 빈 문제에 붙이는 선택. 구독 리스트 클론은 출시는 쉬운데 차별이 어렵다.",
    )

    add_heading_styled(doc, "8. 스택 판단 (이 조사 이후에도 유지)", 2)
    add_bullet(doc, "지금 단계: iOS 네이티브 (SwiftUI) 가정. 혼자이고 아직 제품이 확정되지 않았기 때문.")
    add_bullet(
        doc,
        "크로스플랫폼을 나중에 볼 조건: Android가 꼭 필요한 이유가 생겼을 때. “있으면 좋지”는 이유가 아님.",
    )
    add_bullet(
        doc,
        "네이티브를 고집할 조건: 위젯, 알림, App Intents가 제품의 핵심일 때. 추천안 v1은 여기에 해당한다.",
    )

    add_heading_styled(doc, "9. 참고 링크", 2)
    make_link_table(
        doc,
        [
            (
                "Subo (한국 앱스토어)",
                "https://apps.apple.com/kr/app/구독-관리-갱신-알림-트래커-subo/id6741823650",
            ),
            ("Subo 공식", "https://www.getsubo.app/"),
            ("SubTrack'd", "https://subtrackd.app/"),
            (
                "Subcut 앱스토어",
                "https://apps.apple.com/us/app/subcut-subscription-tracker/id6758765733",
            ),
            (
                "SubManager",
                "https://apps.apple.com/us/app/submanager-subscription-list/id1632853914",
            ),
            (
                "SubTracker (한국 앱스토어)",
                "https://apps.apple.com/kr/app/subtracker-구독-관리/id6759149326",
            ),
            (
                "ReSubs 2026 비교",
                "https://resubs.app/resources/best-subscription-tracker-apps",
            ),
            (
                "CNBC Select 구독 트래커",
                "https://www.cnbc.com/select/best-subscription-trackers/",
            ),
        ],
        [5.0, 11.6],
    )
    add_body(
        doc,
        "가격·평점·기능은 스토어 지역과 시점에 따라 달라질 수 있다. 이 문서는 2026년 8월 15일 웹·스토어 페이지 기준으로 정리했다.",
        before=8,
    )

    doc.save(OUT_PATH)


if __name__ == "__main__":
    build()
    print(OUT_PATH)
